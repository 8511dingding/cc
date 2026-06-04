"""
舆情分析系统 - 主应用 v3 (高级科技风格UI)
"""
import streamlit as st
import pandas as pd
import numpy as np
import json
import re
import os
import hashlib
from copy import copy
from io import BytesIO
from datetime import datetime
from models import (
    init_db, get_session, User, Project, RawData, DataSubset,
    Brand, CompetitorProduct, LabelRule, CleanRule, ReportTemplate, ExportRecord
)
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import plotly.express as px
import plotly.graph_objects as go

# ============== 页面配置 ==============
st.set_page_config(
    page_title="舆情分析系统",
    layout="wide",
    page_icon="📊",
    initial_sidebar_state="expanded"
)

# ============== 浅色专业企业风格 CSS ==============
st.markdown("""
<style>
    /* 全局字体 */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');

    * {
        font-family: 'Inter', 'PingFang SC', 'Microsoft YaHei', sans-serif;
    }

    /* 浅色主题背景 */
    .stApp {
        background: #F5F7FA;
        color: #333333;
    }

    /* 侧边栏 - 深蓝色 */
    [data-testid="stSidebar"] {
        background: #1A3A5C;
        border-right: none;
    }

    [data-testid="stSidebar"] [data-testid="stMarkdownContainer"] {
        color: white;
    }

    /* 顶部栏隐藏 */
    [data-testid="stHeader"] {
        display: none;
    }

    /* 隐藏默认元素 */
    #MainMenu { visibility: hidden; }
    footer { visibility: hidden; }

    /* ========== 通用卡片 ========== */
    .card {
        background: #FFFFFF;
        border-radius: 8px;
        padding: 20px;
        margin: 12px 0;
        box-shadow: 0 1px 3px rgba(0,0,0,0.08);
        border: 1px solid #E8E8E8;
    }

    .card-title {
        font-size: 16px;
        font-weight: 600;
        color: #1A3A5C;
        margin-bottom: 16px;
        padding-bottom: 12px;
        border-bottom: 2px solid #2E7DCD;
    }

    .workflow-hero {
        background: linear-gradient(135deg, #FFFFFF 0%, #F4F8FF 100%);
        border: 1px solid #DDE8F7;
        border-radius: 18px;
        padding: 28px 30px;
        margin: 8px 0 22px;
        box-shadow: 0 16px 40px rgba(26, 58, 92, 0.08);
    }

    .workflow-hero h1 {
        font-size: 30px;
        margin: 0 0 10px;
        letter-spacing: 0;
    }

    .workflow-hero p {
        color: #526071;
        line-height: 1.7;
        margin: 0;
        max-width: 980px;
    }

    .workflow-steps {
        display: grid;
        grid-template-columns: repeat(5, minmax(0, 1fr));
        gap: 12px;
        margin: 18px 0 4px;
    }

    .workflow-step {
        background: rgba(255,255,255,0.86);
        border: 1px solid #E2EAF4;
        border-radius: 12px;
        padding: 14px;
    }

    .workflow-step .num {
        width: 26px;
        height: 26px;
        display: inline-flex;
        align-items: center;
        justify-content: center;
        border-radius: 50%;
        background: #1A73E8;
        color: #FFFFFF;
        font-size: 12px;
        font-weight: 700;
        margin-bottom: 10px;
    }

    .workflow-step .title {
        color: #1A3A5C;
        font-weight: 700;
        font-size: 14px;
        margin-bottom: 4px;
    }

    .workflow-step .desc {
        color: #667085;
        font-size: 12px;
        line-height: 1.5;
    }

    .report-card {
        background: #FFFFFF;
        border: 1px solid #E5EAF1;
        border-radius: 14px;
        padding: 22px;
        margin: 14px 0;
        box-shadow: 0 10px 28px rgba(26, 58, 92, 0.06);
    }

    .report-card h3 {
        margin: 0 0 14px;
        color: #1A3A5C;
        font-size: 17px;
    }

    .quiet-note {
        color: #667085;
        font-size: 13px;
        line-height: 1.7;
    }

    .status-pill {
        display: inline-flex;
        align-items: center;
        gap: 6px;
        border-radius: 999px;
        padding: 5px 11px;
        background: #EAF3FF;
        color: #1A73E8;
        font-size: 12px;
        font-weight: 600;
    }

    @media (max-width: 1100px) {
        .workflow-steps {
            grid-template-columns: repeat(2, minmax(0, 1fr));
        }
    }

    /* ========== 指标卡片 ========== */
    .metric-card {
        background: #FFFFFF;
        border-radius: 8px;
        padding: 20px;
        text-align: center;
        box-shadow: 0 1px 3px rgba(0,0,0,0.08);
        border: 1px solid #E8E8E8;
        transition: all 0.2s;
    }

    .metric-card:hover {
        box-shadow: 0 4px 12px rgba(0,0,0,0.12);
    }

    .metric-card .label {
        color: #666666;
        font-size: 13px;
        font-weight: 500;
        margin-bottom: 8px;
    }

    .metric-card .value {
        color: #2E7DCD;
        font-size: 28px;
        font-weight: 700;
    }

    .metric-card .delta {
        color: #00A5E0;
        font-size: 12px;
        margin-top: 4px;
    }

    /* ========== 标题样式 ========== */
    h1, h2, h3 {
        color: #1A3A5C;
        font-weight: 600;
    }

    h1 {
        font-size: 24px;
        color: #1A3A5C;
    }

    h2 {
        font-size: 18px;
        margin-top: 24px;
    }

    /* ========== 按钮样式 ========== */
    .stButton > button {
        background: #2E7DCD;
        color: white;
        border: none;
        border-radius: 6px;
        font-weight: 500;
        padding: 8px 20px;
        transition: all 0.2s;
    }

    .stButton > button:hover {
        background: #1A5C9C;
        box-shadow: 0 2px 8px rgba(46,125,205,0.3);
    }

    .stButton > button:focus {
        background: #1A5C9C;
    }

    /* 次要按钮 */
    .stButton > button[kind="secondary"] {
        background: #F5F7FA;
        color: #333333;
        border: 1px solid #DDDDDD;
    }

    /* ========== 输入框样式 ========== */
    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea,
    .stNumberInput > div > div > input {
        background: #FFFFFF;
        border: 1px solid #DDDDDD;
        border-radius: 6px;
        color: #333333;
        padding: 10px 14px;
    }

    .stTextInput > div > div > input:focus,
    .stTextArea > div > div > textarea:focus {
        border-color: #2E7DCD;
        box-shadow: 0 0 0 2px rgba(46,125,205,0.15);
    }

    /* ========== 选择框样式 ========== */
    .stSelectbox > div > div {
        background: #FFFFFF;
        border: 1px solid #DDDDDD;
        border-radius: 6px;
    }

    /* ========== 表格样式 ========== */
    [data-testid="stDataFrame"] {
        background: #FFFFFF;
        border-radius: 8px;
        border: 1px solid #E8E8E8;
    }

    /* ========== Tab 样式 ========== */
    .stTabs [data-baseweb="tab-list"] {
        background: #FFFFFF;
        border-radius: 8px;
        padding: 4px;
        border: 1px solid #E8E8E8;
    }

    .stTabs [data-baseweb="tab"] {
        color: #666666;
        font-weight: 500;
        border-radius: 6px;
        padding: 10px 20px;
    }

    .stTabs [data-baseweb="tab"]:hover {
        background: #F0F5FA;
    }

    .stTabs [data-baseweb="tab"][aria-selected="true"] {
        background: #2E7DCD;
        color: white;
    }

    /* ========== 展开器样式 ========== */
    .streamlit-expanderHeader {
        background: #FFFFFF;
        border: 1px solid #E8E8E8;
        border-radius: 8px;
        color: #333333;
        padding: 14px 16px;
    }

    .streamlit-expanderHeader:hover {
        background: #F5F7FA;
    }

    /* ========== 分割线 ========== */
    hr {
        border: none;
        height: 1px;
        background: #E8E8E8;
        margin: 24px 0;
    }

    /* ========== 图表容器 ========== */
    [data-testid="stVegaLiteChart"],
    [data-testid="stPlotlyChart"] {
        background: #FFFFFF;
        border-radius: 8px;
        border: 1px solid #E8E8E8;
        padding: 16px;
    }

    /* ========== 标签样式 ========== */
    .tag {
        display: inline-block;
        padding: 4px 10px;
        border-radius: 4px;
        font-size: 12px;
        font-weight: 500;
    }

    .tag-success {
        background: #E8F5E9;
        color: #2E7D32;
        border: 1px solid #C8E6C9;
    }

    .tag-warning {
        background: #FFF8E1;
        color: #F57C00;
        border: 1px solid #FFE082;
    }

    .tag-info {
        background: #E3F2FD;
        color: #2E7DCD;
        border: 1px solid #BBDEFB;
    }

    .tag-danger {
        background: #FFEBEE;
        color: #D32F2F;
        border: 1px solid #FFCDD2;
    }

    /* ========== SaaS落地页 / 登录入口 ========== */
    .login-page {
        min-height: 100vh;
        background: #FFFFFF;
        color: #142033;
        font-family: 'Inter', 'PingFang SC', 'Microsoft YaHei', sans-serif;
        padding: 0 0 48px;
    }

    .landing-shell {
        width: min(1180px, calc(100vw - 48px));
        margin: 0 auto;
    }

    .landing-nav {
        display: flex;
        align-items: center;
        justify-content: center;
        gap: 28px;
        padding: 24px 0 18px;
        font-size: 14px;
        color: #526071;
    }

    .landing-brand {
        display: flex;
        align-items: center;
        gap: 10px;
        margin-right: auto;
        color: #142033;
        font-size: 16px;
        font-weight: 750;
    }

    .landing-mark {
        width: 32px;
        height: 32px;
        border-radius: 8px;
        background: #1A73E8;
        color: #FFFFFF;
        display: inline-flex;
        align-items: center;
        justify-content: center;
        box-shadow: 0 10px 24px rgba(26, 115, 232, 0.24);
    }

    .landing-nav a {
        color: #526071;
        text-decoration: none;
        font-weight: 600;
    }

    .landing-nav-cta {
        margin-left: auto;
        background: #142033;
        color: #FFFFFF !important;
        padding: 10px 16px;
        border-radius: 8px;
    }

    .landing-hero {
        display: grid;
        grid-template-columns: minmax(0, 0.92fr) minmax(420px, 1.08fr);
        gap: 42px;
        align-items: center;
        padding: 44px 0 64px;
    }

    .landing-hero-copy h1 {
        color: #10203A;
        font-size: clamp(40px, 5vw, 64px);
        line-height: 1.04;
        letter-spacing: 0;
        margin: 0 0 22px;
        max-width: 700px;
    }

    .landing-hero-copy p {
        color: #526071;
        font-size: 17px;
        line-height: 1.85;
        margin: 0 0 28px;
        max-width: 620px;
    }

    .landing-actions {
        display: flex;
        gap: 12px;
        flex-wrap: wrap;
        align-items: center;
    }

    .landing-button {
        display: inline-flex;
        align-items: center;
        justify-content: center;
        min-height: 46px;
        padding: 0 18px;
        border-radius: 8px;
        font-size: 14px;
        font-weight: 750;
        text-decoration: none;
        border: 1px solid #D9E3F0;
    }

    .landing-button.primary {
        background: #1A73E8;
        border-color: #1A73E8;
        color: #FFFFFF;
        box-shadow: 0 14px 26px rgba(26, 115, 232, 0.22);
    }

    .landing-button.secondary {
        background: #FFFFFF;
        color: #142033;
    }

    .product-stage {
        background: linear-gradient(135deg, #F5F9FF 0%, #FFFFFF 56%, #EDF7F4 100%);
        border: 1px solid #DDE8F7;
        border-radius: 16px;
        padding: 18px;
        box-shadow: 0 32px 80px rgba(20, 32, 51, 0.11);
    }

    .product-window {
        background: #FFFFFF;
        border: 1px solid #DDE5EF;
        border-radius: 12px;
        overflow: hidden;
    }

    .product-topbar {
        display: flex;
        align-items: center;
        justify-content: space-between;
        padding: 14px 16px;
        border-bottom: 1px solid #E7EDF5;
        color: #142033;
        font-size: 13px;
        font-weight: 750;
    }

    .window-dots {
        display: flex;
        gap: 6px;
    }

    .window-dot {
        width: 8px;
        height: 8px;
        border-radius: 50%;
        background: #D5DEEA;
    }

    .product-grid {
        display: grid;
        grid-template-columns: 1fr 1.2fr;
        gap: 14px;
        padding: 16px;
    }

    .preview-panel,
    .mini-panel,
    .login-card,
    .workflow-board,
    .report-preview {
        background: #FFFFFF;
        border: 1px solid #E1E8F0;
        border-radius: 8px;
    }

    .preview-panel {
        padding: 14px;
    }

    .panel-label {
        color: #64748B;
        font-size: 12px;
        font-weight: 700;
        margin-bottom: 8px;
    }

    .stat-row {
        display: flex;
        justify-content: space-between;
        gap: 12px;
        padding: 10px 0;
        border-bottom: 1px solid #EEF3F8;
        color: #334155;
        font-size: 12px;
    }

    .stat-row b {
        color: #10203A;
        font-size: 18px;
    }

    .bar-list {
        display: grid;
        gap: 8px;
        margin-top: 12px;
    }

    .bar-line {
        height: 8px;
        border-radius: 99px;
        background: #E8EEF6;
        overflow: hidden;
    }

    .bar-line span {
        display: block;
        height: 100%;
        border-radius: inherit;
        background: #1A73E8;
    }

    .mock-table {
        width: 100%;
        border-collapse: collapse;
        font-size: 12px;
        color: #475569;
    }

    .mock-table th {
        color: #64748B;
        font-size: 11px;
        text-align: left;
        padding: 9px 8px;
        background: #F7FAFD;
        border-bottom: 1px solid #E8EEF6;
    }

    .mock-table td {
        padding: 10px 8px;
        border-bottom: 1px solid #EEF3F8;
    }

    .label-chip {
        display: inline-flex;
        padding: 3px 7px;
        border-radius: 4px;
        background: #EAF3FF;
        color: #1A73E8;
        font-size: 11px;
        font-weight: 700;
    }

    .access-grid {
        display: grid;
        grid-template-columns: 1fr 360px;
        gap: 24px;
        align-items: stretch;
        padding: 0 0 62px;
    }

    .access-copy {
        background: #F6F9FD;
        border: 1px solid #E1E8F0;
        border-radius: 12px;
        padding: 28px;
    }

    .access-copy h2,
    .landing-section h2 {
        color: #10203A;
        font-size: 32px;
        line-height: 1.18;
        margin: 0 0 12px;
        letter-spacing: 0;
    }

    .access-copy p,
    .landing-section p {
        color: #526071;
        font-size: 15px;
        line-height: 1.75;
        margin: 0;
    }

    .proof-list {
        display: grid;
        grid-template-columns: repeat(3, minmax(0, 1fr));
        gap: 12px;
        margin-top: 24px;
    }

    .proof-item {
        background: #FFFFFF;
        border: 1px solid #E1E8F0;
        border-radius: 8px;
        padding: 14px;
        color: #334155;
        font-size: 13px;
        font-weight: 700;
    }

    .login-card {
        padding: 24px;
        box-shadow: 0 18px 48px rgba(20, 32, 51, 0.10);
    }

    .login-title {
        font-size: 22px;
        font-weight: 750;
        color: #10203A;
        margin-bottom: 6px;
    }

    .login-subtitle {
        font-size: 13px;
        color: #64748B;
        margin-bottom: 18px;
    }

    .login-page [data-testid="stForm"] {
        background: transparent;
        border: none;
        padding: 0;
    }

    .login-page [data-testid="stTextInput"] input {
        border: 1px solid #D9E3F0;
        border-radius: 8px;
        min-height: 44px;
        font-size: 14px;
        color: #10203A;
        background: #FFFFFF;
    }

    .login-page [data-testid="stTextInput"] input:focus {
        border-color: #1A73E8;
        box-shadow: 0 0 0 3px rgba(26, 115, 232, 0.12);
    }

    .login-page [data-testid="stFormSubmitButton"] button {
        width: 100%;
        min-height: 44px;
        border-radius: 8px;
        background: #1A73E8;
        border: 1px solid #1A73E8;
        color: #FFFFFF;
        font-size: 14px;
        font-weight: 750;
    }

    .login-footer {
        margin-top: 14px;
        color: #64748B;
        font-size: 12px;
        line-height: 1.6;
    }

    .login-error {
        width: min(1180px, calc(100vw - 48px));
        margin: 0 auto 16px;
        background: #FEF2F2;
        color: #B42318;
        padding: 12px 16px;
        border: 1px solid #FECACA;
        border-radius: 8px;
        font-size: 14px;
    }

    .landing-section {
        padding: 76px 0;
        border-top: 1px solid #EEF3F8;
    }

    .landing-section.blue {
        background: #F6F9FD;
        margin-left: calc(50% - 50vw);
        margin-right: calc(50% - 50vw);
        padding-left: calc(50vw - 50%);
        padding-right: calc(50vw - 50%);
    }

    .workflow-rail {
        display: grid;
        grid-template-columns: repeat(5, minmax(0, 1fr));
        gap: 12px;
        margin: 30px 0 24px;
    }

    .workflow-step-card {
        border-left: 3px solid #1A73E8;
        padding: 8px 10px 8px 12px;
        background: #FFFFFF;
        border-radius: 8px;
        color: #10203A;
        font-size: 13px;
        font-weight: 750;
    }

    .workflow-step-card span {
        display: block;
        color: #64748B;
        font-size: 12px;
        font-weight: 650;
        margin-bottom: 5px;
    }

    .workflow-board {
        display: grid;
        grid-template-columns: 1.4fr 0.8fr;
        gap: 18px;
        padding: 18px;
        box-shadow: 0 18px 48px rgba(20, 32, 51, 0.07);
    }

    .rule-toggle {
        display: flex;
        justify-content: space-between;
        align-items: center;
        padding: 10px 0;
        border-bottom: 1px solid #EEF3F8;
        color: #334155;
        font-size: 13px;
    }

    .toggle-on {
        width: 34px;
        height: 20px;
        border-radius: 99px;
        background: #1A73E8;
        position: relative;
    }

    .toggle-on:after {
        content: "";
        position: absolute;
        width: 14px;
        height: 14px;
        right: 3px;
        top: 3px;
        background: #FFFFFF;
        border-radius: 50%;
    }

    .reports-grid {
        display: grid;
        grid-template-columns: 0.85fr 1.15fr;
        gap: 44px;
        align-items: center;
    }

    .proof-points {
        display: grid;
        gap: 10px;
        margin-top: 22px;
    }

    .proof-point {
        border-bottom: 1px solid #E8EEF6;
        padding: 0 0 12px;
        color: #10203A;
        font-weight: 750;
    }

    .report-preview {
        padding: 18px;
        box-shadow: 0 22px 56px rgba(20, 32, 51, 0.09);
    }

    .report-pages {
        display: grid;
        grid-template-columns: 0.9fr 1.1fr;
        gap: 14px;
        align-items: start;
    }

    .paper {
        background: #FFFFFF;
        border: 1px solid #E1E8F0;
        border-radius: 8px;
        min-height: 220px;
        padding: 18px;
        box-shadow: 0 12px 24px rgba(20, 32, 51, 0.06);
    }

    .paper h4 {
        color: #10203A;
        font-size: 15px;
        margin: 0 0 14px;
    }

    .paper-line {
        height: 8px;
        border-radius: 99px;
        background: #E8EEF6;
        margin-bottom: 10px;
    }

    .landing-cta {
        background: #10203A;
        color: #FFFFFF;
        border-radius: 12px;
        padding: 46px;
    }

    .landing-cta h2 {
        color: #FFFFFF;
        font-size: 34px;
        margin: 0 0 12px;
        max-width: 720px;
    }

    .landing-cta p {
        color: #B8C5D6;
        max-width: 620px;
    }

    .footer-cols {
        display: grid;
        grid-template-columns: repeat(3, 1fr);
        gap: 18px;
        margin-top: 34px;
        padding-top: 28px;
        border-top: 1px solid rgba(255,255,255,0.12);
        color: #B8C5D6;
        font-size: 13px;
        line-height: 1.9;
    }

    .footer-cols b {
        display: block;
        color: #FFFFFF;
        margin-bottom: 8px;
    }

    @media (max-width: 980px) {
        .landing-nav {
            gap: 14px;
            flex-wrap: wrap;
        }

        .landing-brand,
        .landing-nav-cta {
            margin: 0;
        }

        .landing-hero,
        .access-grid,
        .workflow-board,
        .reports-grid,
        .report-pages {
            grid-template-columns: 1fr;
        }

        .product-grid,
        .workflow-rail,
        .proof-list,
        .footer-cols {
            grid-template-columns: 1fr;
        }

        .landing-shell {
            width: min(100% - 28px, 1180px);
        }

        .landing-hero {
            padding-top: 24px;
        }

        .landing-hero-copy h1 {
            font-size: 38px;
        }

        .landing-cta {
            padding: 28px;
        }
    }

    /* ========== 侧边栏导航 ========== */
    .sidebar-logo {
        padding: 20px 16px;
        border-bottom: 1px solid rgba(0,0,0,0.05);
        margin-bottom: 16px;
    }

    .sidebar-logo h2 {
        color: white;
        font-size: 18px;
        margin: 0;
    }

    .sidebar-logo p {
        color: #666666;
        font-size: 11px;
        margin-top: 4px;
    }

    .sidebar-nav {
        padding: 0 8px;
    }

    .nav-item {
        padding: 12px 16px;
        border-radius: 6px;
        color: #444444;
        cursor: pointer;
        transition: all 0.2s;
        margin: 4px 0;
        font-size: 14px;
    }

    .nav-item:hover {
        background: rgba(0,0,0,0.05);
        color: white;
    }

    .nav-item.active {
        background: #2E7DCD;
        color: white;
    }

    .sidebar-user {
        position: absolute;
        bottom: 0;
        left: 0;
        right: 0;
        padding: 16px;
        background: rgba(0,0,0,0.2);
        border-top: 1px solid rgba(0,0,0,0.05);
    }

    .sidebar-user .user-name {
        color: white;
        font-weight: 600;
        font-size: 14px;
    }

    .sidebar-user .user-role {
        color: #666666;
        font-size: 11px;
    }

    /* ========== 进度条 ========== */
    .progress-bar {
        background: #E8E8E8;
        border-radius: 4px;
        height: 8px;
        overflow: hidden;
    }

    .progress-bar .fill {
        background: linear-gradient(90deg, #2E7DCD, #00A5E0);
        height: 100%;
        border-radius: 4px;
        transition: width 0.5s;
    }

    /* ========== 区块标题 ========== */
    .section-title {
        display: flex;
        align-items: center;
        margin: 24px 0 16px;
        padding-bottom: 12px;
        border-bottom: 2px solid #2E7DCD;
    }

    .section-title .icon {
        font-size: 18px;
        margin-right: 10px;
    }

    .section-title h3 {
        margin: 0;
        color: #1A3A5C;
        font-size: 16px;
        font-weight: 600;
    }

    /* ========== 玻璃态效果 ========== */
    .glass {
        background: rgba(255,255,255,0.9);
        backdrop-filter: blur(10px);
        border: 1px solid #E8E8E8;
        border-radius: 8px;
    }

    /* ========== 数据表格 ========== */
    .data-table {
        background: #FFFFFF;
        border-radius: 8px;
        overflow: hidden;
        border: 1px solid #E8E8E8;
    }

    .data-table table {
        width: 100%;
        border-collapse: collapse;
    }

    .data-table th {
        background: #F5F7FA;
        color: #1A3A5C;
        font-weight: 600;
        font-size: 13px;
        padding: 12px 16px;
        text-align: left;
        border-bottom: 2px solid #E8E8E8;
    }

    .data-table td {
        padding: 12px 16px;
        border-bottom: 1px solid #F0F0F0;
        font-size: 14px;
        color: #333333;
    }

    .data-table tr:hover {
        background: #F5F7FA;
    }

    .data-table tr:nth-child(even) {
        background: #FAFAFA;
    }

    .data-table tr:nth-child(even):hover {
        background: #F5F7FA;
    }

    /* ========== 自定义滚动条 ========== */
    ::-webkit-scrollbar {
        width: 6px;
        height: 6px;
    }

    ::-webkit-scrollbar-track {
        background: #F5F7FA;
    }

    ::-webkit-scrollbar-thumb {
        background: #CCCCCC;
        border-radius: 3px;
    }

    ::-webkit-scrollbar-thumb:hover {
        background: #999999;
    }

    /* ========== 分页 ========== */
    .pagination {
        display: flex;
        align-items: center;
        justify-content: center;
        gap: 8px;
        margin-top: 20px;
    }

    /* ========== 状态徽章 ========== */
    .status-badge {
        display: inline-flex;
        align-items: center;
        padding: 4px 10px;
        border-radius: 20px;
        font-size: 12px;
        font-weight: 500;
    }

    .status-badge.online {
        background: #E8F5E9;
        color: #2E7D32;
    }

    .status-badge.offline {
        background: #FFEBEE;
        color: #D32F2F;
    }

    /* ========== 信息提示 ========== */
    .stAlert {
        border-radius: 8px;
    }
</style>
""", unsafe_allow_html=True)

# ============== 辅助函数 ==============

def hash_password(pwd):
    return hashlib.sha256(pwd.encode()).hexdigest()

def verify_password(pwd, hashed):
    return hashlib.sha256(pwd.encode()).hexdigest() == hashed

def is_pure_at(content):
    if pd.isna(content) or content is None:
        return False
    content = str(content)
    if '@' not in content:
        return False
    if content.count('@') > 1:
        return True
    at_pos = content.find('@')
    before = content[:at_pos]
    if before.strip() != '':
        return False
    after = content[at_pos+1:]
    emoji_pattern = re.compile("["
        u"\U0001F600-\U0001F64F"
        u"\U0001F300-\U0001F5FF"
        u"\U0001F680-\U0001F6FF"
        u"\U0001F1E0-\U0001F1FF"
        u"\U00002702-\U000027B0"
        u"\U000024C2-\U0001F251"
        "]+", flags=re.UNICODE)
    cleaned = emoji_pattern.sub('', after)
    cleaned = cleaned.strip(' \t\n\r.,!?;:，。！？；：""''（）()【】[]{}—...·~')
    return cleaned == ''

def is_pure_emoji(content):
    if pd.isna(content) or content is None:
        return False
    content = str(content).strip()
    if content == '':
        return False
    emoji_pattern = re.compile("["
        u"\U0001F600-\U0001F64F"
        u"\U0001F300-\U0001F5FF"
        u"\U0001F680-\U0001F6FF"
        u"\U0001F1E0-\U0001F1FF"
        u"\U00002702-\U000027B0"
        u"\U000024C2-\U0001F251"
        "]+", flags=re.UNICODE)
    cleaned = emoji_pattern.sub('', content)
    cleaned = cleaned.strip(' \t\n\r.,!?;:，。！？；：""''（）()【】[]{}—...·~')
    return cleaned == '' and len(re.findall(emoji_pattern, content)) > 0

def is_filler_word(content):
    if pd.isna(content) or content is None:
        return False
    content = str(content).strip()
    if content == '':
        return False
    emoji_pattern = re.compile("["
        u"\U0001F600-\U0001F64F"
        u"\U0001F300-\U0001F5FF"
        u"\U0001F680-\U0001F6FF"
        u"\U0001F1E0-\U0001F1FF"
        u"\U00002702-\U000027B0"
        u"\U000024C2-\U0001F251"
        "]+", flags=re.UNICODE)
    cleaned = emoji_pattern.sub('', content)
    cleaned = cleaned.strip(' \t\n\r.,!?;:，。！？；：""''（）()【】[]{}—...·~')
    if cleaned == '':
        return False
    filler_words = ['啊', '哈', '嗯', '哦', '呀', '哟', '嘻', '呵', '嘿嘿', '哈哈', '呵呵']
    count = sum(1 for c in cleaned if c in filler_words)
    if len(cleaned) > 0 and count / len(cleaned) > 0.6 and len(cleaned) <= 20:
        return True
    return False

def is_empty_content(content):
    if pd.isna(content) or content is None:
        return True
    content = str(content).strip()
    return content == '' or content == 'nan' or content == 'NaN'

def is_garbled(content):
    if pd.isna(content) or content is None:
        return False
    content = str(content)
    garbled_patterns = [r'[\x00-\x08\x0b\x0c\x0e-\x1f]', r'[▽▃◣◢◇◆]']
    for pattern in garbled_patterns:
        if re.search(pattern, content):
            return True
    weird_chars = sum(1 for c in content if ord(c) > 0xFFFF or ord(c) < 32)
    if len(content) > 5 and weird_chars / len(content) > 0.3:
        return True
    return False

def fuzzy_match(text, keyword):
    if not keyword:
        return True
    text = str(text).lower()
    keyword = keyword.lower()
    if keyword in text:
        return True
    for kw in keyword.split():
        if kw and kw in text:
            return True
    return False

def render_metric_card(label, value, delta=None, icon="📊"):
    """渲染浅色专业风格指标卡片"""
    st.markdown(f"""
    <div class="metric-card">
        <div class="label">{icon} {label}</div>
        <div class="value">{value}</div>
        {f'<div class="delta">{delta}</div>' if delta else ''}
    </div>
    """, unsafe_allow_html=True)

def render_section_title(title, icon="📋"):
    """渲染区块标题"""
    st.markdown(f"""
    <div class="section-title">
        <span class="icon">{icon}</span>
        <h3>{title}</h3>
    </div>
    """, unsafe_allow_html=True)

def render_page_header(title, description="", icon="📊"):
    """渲染页面标题和顶部导航"""
    st.markdown(f"""
    <div style="margin-bottom: 24px;">
        <h1 style="color: #1A3A5C; margin-bottom: 8px;">{icon} {title}</h1>
        {f'<p style="color: #666666; margin: 0;">{description}</p>' if description else ''}
    </div>
    """, unsafe_allow_html=True)

def render_breadcrumb_nav(items):
    """渲染面包屑导航 items: [(name, url_or_page), ...]"""
    nav_html = '<div style="display: flex; align-items: center; gap: 8px; margin-bottom: 16px; font-size: 13px;">'
    for i, (name, _) in enumerate(items):
        if i > 0:
            nav_html += '<span style="color: #999999;">›</span>'
        nav_html += f'<span style="color: {"#2E7DCD" if i == len(items)-1 else "#666666"};">{name}</span>'
    nav_html += '</div>'
    st.markdown(nav_html, unsafe_allow_html=True)

EXCEL_EXPORT_COLUMNS = [
    '视频ID', '视频链接', '关键词', '内容', '话题标签', '发布时间', '博主名',
    '收藏量', '评论量', '点赞量', '分享量', '评论id', '评论时间', 'ip_location',
    '评论内容', '昵称', '二级评论数', '评论获赞', '评论内容类型',
    '认知层阶段一', '认知层阶段二', '情绪层阶段一', '情绪层阶段二',
    '行动层阶段一', '行动层阶段二', '品牌提及'
]

LAYER_LABELS = {
    "cognitive": ["无明确认知", "信息混淆", "精准认知", "泛化抵触"],
    "emotional": ["中性", "正面", "恐慌焦虑", "庆幸旁观", "愤怒背叛"],
    "action": ["暂无行动", "寻求帮助", "转奶流失", "维权诉求"],
}

DEFAULT_LABELS = {
    "cognitive": "无明确认知",
    "emotional": "中性",
    "action": "暂无行动",
}

def safe_text(value):
    if value is None or pd.isna(value):
        return ""
    return str(value)

def to_int(value, default=0):
    try:
        if value is None or pd.isna(value):
            return default
        return int(float(value))
    except (TypeError, ValueError):
        return default

def detect_stage(comment_time):
    text = safe_text(comment_time)
    if text.startswith("2026-04"):
        return "s1"
    if text.startswith("2026-05"):
        return "s2"
    return "s2"

def load_rule_groups(session):
    groups = {}
    rules = session.query(LabelRule).filter(LabelRule.enabled == True).order_by(LabelRule.priority.desc()).all()
    for rule in rules:
        groups.setdefault((rule.layer, rule.stage), []).append(rule)
    return groups

def match_rule_label(text, rules, default_label):
    content = safe_text(text).lower()
    for rule in sorted(rules, key=lambda r: r.priority, reverse=True):
        keywords = rule.keywords or []
        if not keywords:
            continue
        if any(safe_text(keyword).lower() in content for keyword in keywords if safe_text(keyword)):
            return rule.label
    return default_label

def match_brand_labels(text, rules):
    content = safe_text(text).lower()
    matched = []
    for rule in sorted(rules, key=lambda r: r.priority):
        keywords = rule.keywords or []
        if any(safe_text(keyword).lower() in content for keyword in keywords if safe_text(keyword)):
            if rule.label not in matched:
                matched.append(rule.label)
    return "|".join(matched)

def classify_content_type(comment_content, brand_mentions):
    text = safe_text(comment_content)
    if safe_text(brand_mentions):
        return "提及竞品"
    if "@" in text:
        return "@某人互动"
    if len(text) > 100:
        return "长内容(>100字)"
    return "普通内容"

def build_export_dataframe(records):
    rows = []
    for rd in records:
        rows.append({
            '视频ID': rd.video_id or '',
            '视频链接': rd.video_link or '',
            '关键词': rd.keyword or '',
            '内容': rd.content or '',
            '话题标签': rd.topic_tags or '',
            '发布时间': rd.publish_time or '',
            '博主名': rd.blogger or '',
            '收藏量': rd.favorites or 0,
            '评论量': rd.comments_count or 0,
            '点赞量': rd.likes or 0,
            '分享量': rd.shares or 0,
            '评论id': rd.comment_id or '',
            '评论时间': rd.comment_time or '',
            'ip_location': rd.ip_location or '',
            '评论内容': rd.comment_content or '',
            '昵称': rd.nickname or '',
            '二级评论数': rd.reply_count or 0,
            '评论获赞': rd.reply_likes or 0,
            '评论内容类型': rd.content_type or '',
            '认知层阶段一': rd.cognitive_s1 or '',
            '认知层阶段二': rd.cognitive_s2 or '',
            '情绪层阶段一': rd.emotional_s1 or '',
            '情绪层阶段二': rd.emotional_s2 or '',
            '行动层阶段一': rd.action_s1 or '',
            '行动层阶段二': rd.action_s2 or '',
            '品牌提及': rd.brand_detected or rd.brand_mentions or ''
        })
    return pd.DataFrame(rows, columns=EXCEL_EXPORT_COLUMNS)

def dataframe_to_xlsx_bytes(df):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name="舆情数据")
        ws = writer.book["舆情数据"]
        ws.freeze_panes = "A2"
        for cell in ws[1]:
            font = copy(cell.font)
            font.bold = True
            cell.font = font
        widths = {
            'A': 18, 'B': 42, 'C': 16, 'D': 36, 'E': 24, 'F': 20, 'G': 16,
            'H': 10, 'I': 10, 'J': 10, 'K': 10, 'L': 22, 'M': 20, 'N': 12,
            'O': 48, 'P': 16, 'Q': 10, 'R': 10, 'S': 16, 'T': 16, 'U': 16,
            'V': 16, 'W': 16, 'X': 16, 'Y': 16, 'Z': 24
        }
        for col, width in widths.items():
            ws.column_dimensions[col].width = width
    output.seek(0)
    return output.getvalue()

def apply_auto_labels(project_id, overwrite=False):
    session = get_session()
    rule_groups = load_rule_groups(session)
    records = session.query(RawData).filter(RawData.project_id == project_id, RawData.is_valid == True).all()
    updated = 0
    for rd in records:
        text = f"{safe_text(rd.comment_content)} {safe_text(rd.content)}"
        stage = detect_stage(rd.comment_time)
        brand_text = f"{text} {safe_text(rd.brand_mentions)}"
        brand_hit = match_brand_labels(brand_text, rule_groups.get(("brand", stage), []))
        target_fields = {
            "cognitive": f"cognitive_{stage}",
            "emotional": f"emotional_{stage}",
            "action": f"action_{stage}",
        }
        for layer, field in target_fields.items():
            current = getattr(rd, field, None)
            if overwrite or not current:
                setattr(rd, field, match_rule_label(text, rule_groups.get((layer, stage), []), DEFAULT_LABELS[layer]))
                updated += 1
        if overwrite or not rd.brand_detected:
            rd.brand_detected = brand_hit or rd.brand_mentions or ""
            updated += 1
        if not rd.content_type:
            rd.content_type = classify_content_type(rd.comment_content, rd.brand_detected or rd.brand_mentions)
    session.commit()
    session.close()
    return updated, len(records)

def get_project_dataframe(project_id):
    session = get_session()
    records = session.query(RawData).filter(RawData.project_id == project_id).all()
    project = session.query(Project).filter(Project.id == project_id).first()
    data = pd.DataFrame([r.to_dict() for r in records])
    session.close()
    return project, data

def stage_masks(df):
    if len(df) == 0:
        return pd.Series(dtype=bool), pd.Series(dtype=bool)
    s1 = df['comment_time'].astype(str).str.startswith('2026-04')
    s2 = df['comment_time'].astype(str).str.startswith('2026-05')
    return s1, s2

def label_counts(df, column, labels):
    stats = df[column].fillna("").value_counts().to_dict() if column in df.columns else {}
    return [{"标签": label, "数量": int(stats.get(label, 0))} for label in labels]

def value_counts_df(df, column, name_col, top=None, empty_label="未识别"):
    if column not in df.columns:
        return pd.DataFrame(columns=[name_col, "数量"])
    series = df[column].fillna("").replace("", empty_label).value_counts()
    if top:
        series = series.head(top)
    return pd.DataFrame({name_col: series.index.tolist(), "数量": series.values.astype(int).tolist()})

def build_report_context(project, df):
    if len(df) == 0:
        return {"ready": False}
    valid_df = df[df['is_valid'] == True].copy()
    s1_mask, s2_mask = stage_masks(valid_df)
    s1_df = valid_df[s1_mask]
    s2_df = valid_df[s2_mask]
    s1_d = project.s1_denominator or max(len(s1_df), 1)
    s2_d = project.s2_denominator or max(len(s2_df), 1)
    return {
        "ready": True,
        "project_name": project.name if project else "舆情分析项目",
        "total": len(df),
        "valid": len(valid_df),
        "invalid": len(df) - len(valid_df),
        "s1_count": len(s1_df),
        "s2_count": len(s2_df),
        "s1_denominator": s1_d,
        "s2_denominator": s2_d,
        "content_type": value_counts_df(valid_df, "content_type", "评论类型", empty_label="未分类"),
        "cognitive_s1": label_counts(s1_df, "cognitive_s1", LAYER_LABELS["cognitive"]),
        "cognitive_s2": label_counts(s2_df, "cognitive_s2", LAYER_LABELS["cognitive"]),
        "emotional_s1": label_counts(s1_df, "emotional_s1", LAYER_LABELS["emotional"]),
        "emotional_s2": label_counts(s2_df, "emotional_s2", LAYER_LABELS["emotional"]),
        "action_s1": label_counts(s1_df, "action_s1", LAYER_LABELS["action"]),
        "action_s2": label_counts(s2_df, "action_s2", LAYER_LABELS["action"]),
        "brand": value_counts_df(valid_df, "brand_detected", "品牌", top=20, empty_label="无品牌"),
    }

def render_stat_table(title, rows, denom):
    st.markdown(f"<div class='report-card'><h3>{title}</h3>", unsafe_allow_html=True)
    df = pd.DataFrame(rows)
    if len(df) == 0:
        st.info("暂无数据")
    else:
        df["占比"] = df["数量"].apply(lambda x: f"{x / denom * 100:.1f}%" if denom else "0.0%")
        st.dataframe(df, hide_index=True, use_container_width=True, height=min(260, 72 + len(df) * 42))
    st.markdown("</div>", unsafe_allow_html=True)

def add_doc_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = str(header)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    return table

def add_report_section_table(doc, title, rows, denom):
    doc.add_heading(title, level=2)
    table_rows = []
    for item in rows:
        count = item["数量"]
        pct = f"{count / denom * 100:.1f}%" if denom else "0.0%"
        table_rows.append([item["标签"], count, pct])
    add_doc_table(doc, ["类型", "数量", "占比"], table_rows)

def build_word_report(project, df, context):
    doc = Document()
    title = doc.add_heading(f"{context['project_name']} 舆情分析报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("一、报告概述", level=1)
    doc.add_paragraph(f"分析目的：帮助品牌了解消费者如何解读舆情事件，掌握用户情绪、认知变化和行动倾向。")
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"总评论数：{context['total']}条；有效评论数：{context['valid']}条；无效评论数：{context['invalid']}条。")
    doc.add_paragraph(f"第一阶段：{context['s1_count']}条；第二阶段：{context['s2_count']}条。")
    doc.add_heading("二、评论内容类型分布", level=1)
    ctype = context["content_type"]
    if len(ctype) > 0:
        add_doc_table(doc, ["评论类型", "数量"], ctype.astype(str).values.tolist())
    doc.add_heading("三、认知分析（Understanding）", level=1)
    doc.add_paragraph("认知解码层分析，评估消费者对事件说明、产品版本和召回范围的理解程度。")
    add_report_section_table(doc, "【第一阶段】", context["cognitive_s1"], context["s1_denominator"])
    add_report_section_table(doc, "【第二阶段】", context["cognitive_s2"], context["s2_denominator"])
    doc.add_heading("四、情绪分析（Emotional）", level=1)
    doc.add_paragraph("情绪图谱层分析，评估舆论心理状态和安抚优先级。")
    add_report_section_table(doc, "【第一阶段】", context["emotional_s1"], context["s1_denominator"])
    add_report_section_table(doc, "【第二阶段】", context["emotional_s2"], context["s2_denominator"])
    doc.add_heading("五、行为分析（Action）", level=1)
    doc.add_paragraph("行为倾向层分析，识别消费者是否仅表达情绪，或已出现求助、转奶和维权诉求。")
    add_report_section_table(doc, "【第一阶段】", context["action_s1"], context["s1_denominator"])
    add_report_section_table(doc, "【第二阶段】", context["action_s2"], context["s2_denominator"])
    doc.add_heading("六、品牌竞品提及分析", level=1)
    brand = context["brand"]
    if len(brand) > 0:
        add_doc_table(doc, ["品牌", "数量"], brand.astype(str).values.tolist())
    doc.add_heading("七、总结与建议", level=1)
    doc.add_paragraph("建议结合人工校正后的标签继续迭代关键词规则，并对高风险情绪和行动倾向样本进行重点复核。")
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()

def extract_learning_suggestions(project_id, min_count=2):
    session = get_session()
    records = session.query(RawData).filter(
        RawData.project_id == project_id,
        RawData.manual_override == True,
        RawData.is_valid == True
    ).all()
    rules = session.query(LabelRule).all()
    existing = {}
    for rule in rules:
        existing.setdefault((rule.layer, rule.stage, rule.label), set()).update(rule.keywords or [])
    session.close()

    targets = [
        ("cognitive", "s1", "cognitive_s1"),
        ("cognitive", "s2", "cognitive_s2"),
        ("emotional", "s1", "emotional_s1"),
        ("emotional", "s2", "emotional_s2"),
        ("action", "s1", "action_s1"),
        ("action", "s2", "action_s2"),
    ]
    buckets = {}
    for rd in records:
        text = safe_text(rd.comment_content)
        tokens = set(re.findall(r'[\u4e00-\u9fa5]{2,8}|[A-Za-z0-9]{2,12}', text))
        tokens = {t for t in tokens if len(t) >= 2 and t not in {"我们", "这个", "那个", "宝宝", "奶粉", "评论", "真的"}}
        for layer, stage, attr in targets:
            label = getattr(rd, attr, None)
            if not label:
                continue
            key = (layer, stage, label)
            known = existing.get(key, set())
            for token in tokens:
                if token not in known:
                    buckets.setdefault((layer, stage, label, token), 0)
                    buckets[(layer, stage, label, token)] += 1
    rows = []
    for (layer, stage, label, token), count in buckets.items():
        if count >= min_count:
            rows.append({"层级": layer, "阶段": stage, "标签": label, "建议关键词": token, "出现次数": count})
    return pd.DataFrame(rows).sort_values(["出现次数", "层级"], ascending=[False, True]) if rows else pd.DataFrame(columns=["层级", "阶段", "标签", "建议关键词", "出现次数"])

def sync_market_brand_rules():
    rules_path = os.path.join(os.path.dirname(__file__), "brand_rules.json")
    if not os.path.exists(rules_path):
        return
    with open(rules_path, "r", encoding="utf-8") as f:
        brand_rules = json.load(f)
    session = get_session()
    for brand in brand_rules:
        for stage in ["s1", "s2"]:
            rule = session.query(LabelRule).filter(
                LabelRule.layer == "brand",
                LabelRule.stage == stage,
                LabelRule.label == brand["label"]
            ).first()
            if rule:
                merged = list(dict.fromkeys((rule.keywords or []) + brand["keywords"]))
                rule.keywords = merged
                rule.enabled = True
            else:
                session.add(LabelRule(
                    layer="brand",
                    stage=stage,
                    label=brand["label"],
                    keywords=brand["keywords"],
                    priority=brand["priority"],
                    enabled=True
                ))
    session.commit()
    session.close()

# ============== 数据库初始化 ==============
init_db()
sync_market_brand_rules()

# ============== SaaS落地页 / 登录入口 ==============
if 'user_id' not in st.session_state:
    st.markdown("""
    <div class="login-page">
      <div class="landing-shell">
        <nav class="landing-nav" aria-label="主导航">
          <div class="landing-brand">
            <span class="landing-mark">
              <svg width="18" height="18" viewBox="0 0 24 24" fill="none" aria-hidden="true">
                <path d="M5 18V8M12 18V4M19 18v-7" stroke="currentColor" stroke-width="2.2" stroke-linecap="round"/>
              </svg>
            </span>
            <span>舆情标签闭环工作台</span>
          </div>
          <a href="#workflow">工作流</a>
          <a href="#labels">标签体系</a>
          <a href="#reports">报告导出</a>
          <a href="#deployment">安全部署</a>
          <a class="landing-nav-cta" href="#access">预约演示</a>
        </nav>

        <section class="landing-hero">
          <div class="landing-hero-copy">
            <h1>把社媒舆情数据，变成可复核的品牌洞察</h1>
            <p>从 Excel 导入、清洗规则、四层标签到 Word/PDF 报告导出，让团队用同一套口径完成分析交付。</p>
            <div class="landing-actions">
              <a class="landing-button primary" href="#access">预约演示</a>
              <a class="landing-button secondary" href="#workflow">查看工作流</a>
            </div>
          </div>
          <div class="product-stage" aria-label="产品界面预览">
            <div class="product-window">
              <div class="product-topbar">
                <span>项目总览 / a2 召回舆情复盘</span>
                <span class="window-dots"><span class="window-dot"></span><span class="window-dot"></span><span class="window-dot"></span></span>
              </div>
              <div class="product-grid">
                <div class="preview-panel">
                  <div class="panel-label">数据状态</div>
                  <div class="stat-row"><span>总评论数</span><b>20,486</b></div>
                  <div class="stat-row"><span>有效评论</span><b>16,802</b></div>
                  <div class="stat-row"><span>已标注</span><b>14,219</b></div>
                  <div class="bar-list">
                    <div class="bar-line"><span style="width: 82%;"></span></div>
                    <div class="bar-line"><span style="width: 64%; background:#15A585;"></span></div>
                    <div class="bar-line"><span style="width: 48%; background:#F4A62A;"></span></div>
                  </div>
                </div>
                <div class="preview-panel">
                  <div class="panel-label">标签分布与阶段对比</div>
                  <table class="mock-table">
                    <thead><tr><th>层级</th><th>阶段一</th><th>阶段二</th><th>状态</th></tr></thead>
                    <tbody>
                      <tr><td>认知</td><td>断货</td><td>召回</td><td><span class="label-chip">已复核</span></td></tr>
                      <tr><td>情绪</td><td>焦虑</td><td>愤怒</td><td><span class="label-chip">需校正</span></td></tr>
                      <tr><td>行动</td><td>观望</td><td>转奶</td><td><span class="label-chip">报告中</span></td></tr>
                      <tr><td>品牌</td><td>a2</td><td>飞鹤</td><td><span class="label-chip">导出</span></td></tr>
                    </tbody>
                  </table>
                </div>
              </div>
            </div>
          </div>
        </section>

        <section class="access-grid" id="access">
          <div class="access-copy">
            <h2>同一套规则，贯穿清洗、标注、复核和导出</h2>
            <p>品牌、竞品、阶段和标签规则统一沉淀在系统内，减少临时 Excel 往返和报告口径不一致。</p>
            <div class="proof-list">
              <div class="proof-item">四层标签体系</div>
              <div class="proof-item">两阶段并列对比</div>
              <div class="proof-item">Word / PDF 报告</div>
            </div>
          </div>
          <div class="login-card">
            <div class="login-title">打开本地系统</div>
            <div class="login-subtitle">使用已配置账号进入舆情分析工作台</div>
    """, unsafe_allow_html=True)

    with st.form("login_form"):
        username = st.text_input("", placeholder="用户名 / Username", label_visibility="collapsed")
        password = st.text_input("", placeholder="密码 / Password", type="password", label_visibility="collapsed")
        submitted = st.form_submit_button("登录工作台", type="primary")

    st.markdown("""
            <div class="login-footer">默认测试账号：admin / admin123</div>
          </div>
        </section>

        <section class="landing-section blue" id="workflow">
          <div class="landing-shell">
            <h2>从导入到确认，分析过程清晰可追踪</h2>
            <p>每一步都保留可复核的样本、规则和状态，适合多人协作完成交付型舆情分析。</p>
            <div class="workflow-rail">
              <div class="workflow-step-card"><span>01</span>导入数据</div>
              <div class="workflow-step-card"><span>02</span>清洗预览</div>
              <div class="workflow-step-card"><span>03</span>自动打标</div>
              <div class="workflow-step-card"><span>04</span>人工校正</div>
              <div class="workflow-step-card"><span>05</span>报告确认</div>
            </div>
            <div class="workflow-board" id="labels">
              <div>
                <div class="panel-label">样本与标签复核</div>
                <table class="mock-table">
                  <thead><tr><th>评论摘要</th><th>认知</th><th>情绪</th><th>行动</th><th>品牌</th></tr></thead>
                  <tbody>
                    <tr><td>担心召回批次影响宝宝...</td><td>安全疑虑</td><td>焦虑</td><td>观望</td><td>a2</td></tr>
                    <tr><td>已经开始比较其他品牌...</td><td>品牌对比</td><td>失望</td><td>转奶</td><td>飞鹤</td></tr>
                    <tr><td>客服解释后仍希望公布...</td><td>信息透明</td><td>质疑</td><td>维权</td><td>爱他美</td></tr>
                    <tr><td>囤货用户最关心退换...</td><td>售后政策</td><td>焦虑</td><td>咨询</td><td>a2</td></tr>
                  </tbody>
                </table>
              </div>
              <div>
                <div class="panel-label">清洗规则实时预览</div>
                <div class="rule-toggle"><span>纯表情 / emoji</span><span class="toggle-on"></span></div>
                <div class="rule-toggle"><span>无意义语气词</span><span class="toggle-on"></span></div>
                <div class="rule-toggle"><span>重复内容</span><span class="toggle-on"></span></div>
                <div class="rule-toggle"><span>乱码异常字符</span><span class="toggle-on"></span></div>
              </div>
            </div>
          </div>
        </section>

        <section class="landing-section" id="reports">
          <div class="reports-grid">
            <div>
              <h2>交付前，每个结论都能回到原始样本</h2>
              <p>模板、板块、文件名和导出记录都在系统内完成管理，减少报告版本分叉和人工核对成本。</p>
              <div class="proof-points">
                <div class="proof-point">四层标签统计</div>
                <div class="proof-point">两阶段并列对比</div>
                <div class="proof-point">Word / PDF 导出</div>
              </div>
            </div>
            <div class="report-preview">
              <div class="report-pages">
                <div class="paper">
                  <h4>舆情分析报告</h4>
                  <div class="paper-line" style="width: 80%;"></div>
                  <div class="paper-line" style="width: 62%;"></div>
                  <div class="paper-line" style="width: 92%;"></div>
                  <div class="bar-list">
                    <div class="bar-line"><span style="width: 70%;"></span></div>
                    <div class="bar-line"><span style="width: 44%; background:#15A585;"></span></div>
                    <div class="bar-line"><span style="width: 56%; background:#F4A62A;"></span></div>
                  </div>
                </div>
                <div class="preview-panel">
                  <div class="panel-label">导出记录</div>
                  <table class="mock-table">
                    <thead><tr><th>文件名</th><th>格式</th><th>状态</th></tr></thead>
                    <tbody>
                      <tr><td>a2_舆情分析报告</td><td>DOCX</td><td><span class="label-chip">已确认</span></td></tr>
                      <tr><td>竞品标签复核表</td><td>XLSX</td><td><span class="label-chip">可下载</span></td></tr>
                      <tr><td>阶段对比摘要</td><td>PDF</td><td><span class="label-chip">归档</span></td></tr>
                    </tbody>
                  </table>
                </div>
              </div>
            </div>
          </div>
        </section>

        <section class="landing-section" id="deployment">
          <div class="landing-cta">
            <h2>让下一份舆情报告少一点返工，多一点依据</h2>
            <p>为品牌、竞品、阶段和标签规则建立统一工作台，把团队的分析经验沉淀为可复用流程。</p>
            <div class="landing-actions" style="margin-top: 24px;">
              <a class="landing-button primary" href="#access">预约演示</a>
              <a class="landing-button secondary" href="#access">打开本地系统</a>
            </div>
            <div class="footer-cols">
              <div><b>产品</b>数据导入<br>标签管理<br>报告生成</div>
              <div><b>资源</b>清洗规则<br>模板配置<br>导出记录</div>
              <div><b>部署</b>SQLite 本地库<br>Streamlit 服务<br>Nginx 子路径</div>
            </div>
          </div>
        </section>
      </div>
    """, unsafe_allow_html=True)

    # 错误提示
    if 'login_error' in st.session_state:
        st.markdown(f'<div class="login-error">{st.session_state.login_error}</div>', unsafe_allow_html=True)
        del st.session_state['login_error']

    st.markdown('</div>', unsafe_allow_html=True)

    if submitted:
        session = get_session()
        user = session.query(User).filter(User.username == username, User.is_active == True).first()
        session.close()
        if user and verify_password(password, user.password_hash):
            st.session_state['user_id'] = user.id
            st.session_state['username'] = user.username
            st.session_state['display_name'] = user.display_name or user.username
            st.session_state['role'] = user.role
            st.rerun()
        else:
            st.session_state['login_error'] = "用户名或密码错误"
            st.rerun()

    st.stop()

# ============== 侧边栏 ==============
with st.sidebar:
    st.markdown("""
    <div class="sidebar-logo">
        <h2>📊 舆情分析</h2>
        <p>Sentiment Analysis Platform</p>
    </div>
    """, unsafe_allow_html=True)

    session = get_session()
    projects = session.query(Project).filter(Project.status == 'active').all()
    session.close()

    project_names = [p.name for p in projects]
    selected = st.selectbox("", ["📁 选择项目"] + ["➕ 创建新项目"] + project_names, label_visibility="collapsed")

    if selected == "➕ 创建新项目":
        new_name = st.text_input("", placeholder="项目名称", label_visibility="collapsed")
        if st.button("创建", use_container_width=True):
            if new_name:
                session = get_session()
                p = Project(name=new_name, created_by=st.session_state['user_id'])
                session.add(p)
                session.commit()
                session.close()
                st.success("创建成功")
                st.rerun()
    elif selected != "📁 选择项目":
        session = get_session()
        current_project = session.query(Project).filter(Project.name == selected).first()
        session.close()
        if current_project:
            st.session_state['current_project_id'] = current_project.id
            st.session_state['current_project_name'] = current_project.name

    # 如果没有项目，自动创建一个默认项目
    if 'current_project_id' not in st.session_state:
        session = get_session()
        default_project = session.query(Project).filter(Project.status == 'active').first()
        if not default_project:
            default_project = Project(
                name="舆情分析项目",
                description="默认项目",
                created_by=st.session_state.get('user_id')
            )
            session.add(default_project)
            session.commit()
        st.session_state['current_project_id'] = default_project.id
        st.session_state['current_project_name'] = default_project.name
        session.close()

    project_id = st.session_state['current_project_id']

    menu_items = [
        ("🏠", "工作台"),
        ("📥", "数据导入"),
        ("🧹", "数据清洗"),
        ("⚙️", "自动打标"),
        ("✏️", "手动标注"),
        ("🧠", "规则学习"),
        ("📄", "报告生成"),
        ("📊", "数据总览"),
        ("🏷️", "标签管理"),
        ("📊", "数据统计"),
        ("📋", "内容管理"),
        ("📝", "规则管理"),
        ("🏭", "品牌竞品"),
        ("📋", "数据子集"),
        ("📋", "模板管理"),
        ("📁", "导出记录"),
    ]

    if st.session_state.get('role') == 'admin':
        menu_items.append(("👤", "用户管理"))

    menu_items.append(("⚙️", "系统设置"))

    selected_menu = st.radio("", [f"{icon} {name}" for icon, name in menu_items], label_visibility="collapsed")

    menu_map = {f"{icon} {name}": name for icon, name in menu_items}
    page = menu_map.get(selected_menu, "工作台")

    st.divider()
    st.markdown(f"""
    <div class="sidebar-user">
        <div class="user-name">{st.session_state.get('display_name', '用户')}</div>
        <div class="user-role">{st.session_state.get('role', 'user').upper()}</div>
    </div>
    """, unsafe_allow_html=True)

    if st.button("🚪 退出登录", use_container_width=True):
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.rerun()

# ============== 工作台页 ==============
if page == "工作台":
    project, df = get_project_dataframe(project_id)
    report_ok = st.session_state.get(f"report_confirmed_{project_id}", False)
    st.markdown("""
    <div class="workflow-hero">
        <h1>舆情标签闭环工作台</h1>
        <p>从数据导入开始，系统先按规则自动打标；人工在大表格里校正；系统根据校正样本提出规则优化建议；在线报告确认无误后，再输出 Word 报告和参考格式 Excel。</p>
        <div class="workflow-steps">
            <div class="workflow-step"><div class="num">1</div><div class="title">导入数据</div><div class="desc">按客户Excel列顺序导入评论和内容。</div></div>
            <div class="workflow-step"><div class="num">2</div><div class="title">自动打标</div><div class="desc">根据认知/情绪/行动/品牌规则批量匹配。</div></div>
            <div class="workflow-step"><div class="num">3</div><div class="title">人工校正</div><div class="desc">在Excel式表格里复核和修改标签。</div></div>
            <div class="workflow-step"><div class="num">4</div><div class="title">规则学习</div><div class="desc">从人工修改中提取新增关键词建议。</div></div>
            <div class="workflow-step"><div class="num">5</div><div class="title">报告确认</div><div class="desc">在线预览报告，确认后导出Word/Excel。</div></div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    if len(df) == 0:
        st.info("当前项目暂无数据。请先进入「数据导入」上传Excel。")
    else:
        valid_count = int(df['is_valid'].sum())
        manual_count = int(df['manual_override'].fillna(False).sum())
        labeled_count = int(((df['cognitive_s1'].fillna("") != "") | (df['cognitive_s2'].fillna("") != "")).sum())
        label_rate = labeled_count / len(df) * 100 if len(df) else 0
        session = get_session()
        rule_count = session.query(LabelRule).filter(LabelRule.enabled == True).count()
        session.close()

        cols = st.columns(5)
        metrics = [
            ("总数据", f"{len(df):,}", "📦"),
            ("有效数据", f"{valid_count:,}", "✅"),
            ("已打标签", f"{labeled_count:,}", "🏷️"),
            ("人工校正", f"{manual_count:,}", "✏️"),
            ("启用规则", f"{rule_count:,}", "🧠"),
        ]
        for col, item in zip(cols, metrics):
            with col:
                render_metric_card(item[0], item[1], icon=item[2])

        st.markdown("<div class='report-card'>", unsafe_allow_html=True)
        st.markdown("### 当前闭环状态")
        st.progress(min(label_rate / 100, 1.0), text=f"标注完成度 {label_rate:.1f}%")
        status_cols = st.columns(4)
        with status_cols[0]:
            st.markdown("<span class='status-pill'>数据已导入</span>" if len(df) else "<span class='status-pill'>等待导入</span>", unsafe_allow_html=True)
        with status_cols[1]:
            st.markdown("<span class='status-pill'>可自动打标</span>", unsafe_allow_html=True)
        with status_cols[2]:
            st.markdown(f"<span class='status-pill'>{manual_count} 条人工校正</span>", unsafe_allow_html=True)
        with status_cols[3]:
            st.markdown("<span class='status-pill'>报告已确认</span>" if report_ok else "<span class='status-pill'>报告待确认</span>", unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown("<div class='report-card'>", unsafe_allow_html=True)
        st.markdown("### 最近20条样本")
        preview_cols = ['comment_content', 'comment_time', 'content_type', 'cognitive_s2', 'emotional_s2', 'action_s2', 'brand_detected']
        st.dataframe(df[preview_cols].tail(20), hide_index=True, use_container_width=True, height=520)
        st.markdown("</div>", unsafe_allow_html=True)

# ============== 数据导入页 ==============
elif page == "数据导入":
    render_page_header("📥 数据导入", "上传Excel数据文件，系统将自动进行数据清洗和预处理")

    uploaded_file = st.file_uploader("选择Excel文件", type=['xlsx', 'xls'])

    if uploaded_file:
        file_size = len(uploaded_file.getvalue()) / 1024
        st.markdown(f"""
        <div class="card" style="padding: 20px; margin: 16px 0;">
            <div style="display: flex; justify-content: space-between; align-items: center;">
                <div>
                    <div style="font-size: 14px; color: #333333;">📎 {uploaded_file.name}</div>
                    <div style="font-size: 12px; color: #666666; margin-top: 4px;">{file_size:.1f} KB</div>
                </div>
                <div class="tag tag-info">待导入</div>
            </div>
        </div>
        """, unsafe_allow_html=True)

        auto_after_import = st.checkbox("导入完成后立即按当前规则自动打标签", value=True)

        if st.button("🚀 开始导入", use_container_width=True):
            with st.spinner("正在导入数据..."):
                df = pd.read_excel(uploaded_file)

                # 分批次导入，每批1万条，支持10万级数据
                batch_size = 10000
                total_rows = len(df)

                # 先删除旧数据
                session = get_session()
                session.query(RawData).filter(RawData.project_id == project_id).delete()
                session.commit()
                session.close()

                # 批量插入新数据
                for batch_start in range(0, total_rows, batch_size):
                    batch_end = min(batch_start + batch_size, total_rows)
                    session = get_session()

                    for idx in range(batch_start, batch_end):
                        row = df.iloc[idx]
                        comment_content = str(row.get('评论内容', '')) if pd.notna(row.get('评论内容')) else ''
                        rd = RawData(
                            project_id=project_id,
                            row_index=int(idx) if pd.notna(idx) else idx,
                            video_id=str(row.get('视频ID', '')) if pd.notna(row.get('视频ID')) else '',
                            video_link=str(row.get('视频链接', '')) if pd.notna(row.get('视频链接')) else '',
                            keyword=str(row.get('关键词', '')) if pd.notna(row.get('关键词')) else '',
                            content=str(row.get('内容', '')) if pd.notna(row.get('内容')) else '',
                            topic_tags=str(row.get('话题标签', '')) if pd.notna(row.get('话题标签')) else '',
                            publish_time=str(row.get('发布时间', '')) if pd.notna(row.get('发布时间')) else '',
                            blogger=str(row.get('博主名', '')) if pd.notna(row.get('博主名')) else '',
                            likes=int(row.get('点赞量', 0)) if pd.notna(row.get('点赞量')) else 0,
                            comments_count=int(row.get('评论量', 0)) if pd.notna(row.get('评论量')) else 0,
                            favorites=int(row.get('收藏量', 0)) if pd.notna(row.get('收藏量')) else 0,
                            shares=int(row.get('分享量', 0)) if pd.notna(row.get('分享量')) else 0,
                            comment_id=str(row.get('评论id', '')) if pd.notna(row.get('评论id')) else '',
                            comment_time=str(row.get('评论时间', '')) if pd.notna(row.get('评论时间')) else '',
                            ip_location=str(row.get('ip_location', '')) if pd.notna(row.get('ip_location')) else '',
                            comment_content=comment_content,
                            nickname=str(row.get('昵称', '')) if pd.notna(row.get('昵称')) else '',
                            reply_count=int(row.get('二级评论数', 0)) if pd.notna(row.get('二级评论数')) else 0,
                            reply_likes=int(row.get('评论获赞', 0)) if pd.notna(row.get('评论获赞')) else 0,
                            content_type=str(row.get('评论内容类型', '')) if pd.notna(row.get('评论内容类型')) else classify_content_type(comment_content, str(row.get('品牌提及', '')) if pd.notna(row.get('品牌提及')) else ''),
                            brand_mentions=str(row.get('品牌提及', '')) if pd.notna(row.get('品牌提及')) else '',
                            cognitive_s1=str(row.get('认知层阶段一', '')) if pd.notna(row.get('认知层阶段一')) else '',
                            cognitive_s2=str(row.get('认知层阶段二', '')) if pd.notna(row.get('认知层阶段二')) else '',
                            emotional_s1=str(row.get('情绪层阶段一', '')) if pd.notna(row.get('情绪层阶段一')) else '',
                            emotional_s2=str(row.get('情绪层阶段二', '')) if pd.notna(row.get('情绪层阶段二')) else '',
                            action_s1=str(row.get('行动层阶段一', '')) if pd.notna(row.get('行动层阶段一')) else '',
                            action_s2=str(row.get('行动层阶段二', '')) if pd.notna(row.get('行动层阶段二')) else '',
                        )
                        rd.is_pure_at = is_pure_at(comment_content)
                        rd.is_pure_emoji = is_pure_emoji(comment_content)
                        rd.is_filler_word = is_filler_word(comment_content)
                        rd.is_empty = is_empty_content(comment_content)
                        rd.is_garbled = is_garbled(comment_content)
                        rd.is_valid = not any([rd.is_pure_at, rd.is_pure_emoji, rd.is_filler_word, rd.is_empty, rd.is_garbled])
                        session.add(rd)

                    session.commit()
                    session.close()
                    progress = min(batch_end, total_rows)
                    st.info(f"已导入 {progress}/{total_rows} 条...")

                # 更新项目统计
                session = get_session()
                project = session.query(Project).filter(Project.id == project_id).first()
                if project:
                    project.total_rows = total_rows
                    project.source_file = uploaded_file.name
                    valid_count = session.query(RawData).filter(RawData.project_id == project_id, RawData.is_valid == True).count()
                    project.valid_rows = valid_count
                    session.commit()
                session.close()

                if auto_after_import:
                    updated, valid_records = apply_auto_labels(project_id, overwrite=False)
                    st.info(f"自动打标完成：有效数据 {valid_records} 条，填充字段 {updated} 个。")

            st.success(f"✅ 成功导入 {total_rows} 条数据！")
            st.rerun()

    session = get_session()
    count = session.query(RawData).filter(RawData.project_id == project_id).count()
    session.close()

    if count > 0:
        st.markdown("<hr>", unsafe_allow_html=True)
        render_section_title("📊 数据概览", "📊")

        session = get_session()
        records = session.query(RawData).filter(RawData.project_id == project_id).all()
        df_all = pd.DataFrame([r.to_dict() for r in records])
        session.close()

        total = len(df_all)
        valid = df_all['is_valid'].sum()
        s1 = df_all['comment_time'].astype(str).str.startswith('2026-04').sum()

        cols = st.columns(4)
        for col, (label, value, icon) in zip(cols, [
            ("总数据量", f"{total} 条", "📦"),
            ("有效数据", f"{valid} 条", "✅"),
            ("已清洗", f"{total - valid} 条", "🧹"),
            ("第一阶段", f"{s1} 条", "🗓️")
        ]):
            with col:
                render_metric_card(label, value, icon=icon)

# ============== 数据总览页 ==============
elif page == "数据总览":
    render_page_header("📊 数据总览", "查看项目数据总览和统计信息")

    session = get_session()
    records = session.query(RawData).filter(RawData.project_id == project_id).all()
    df = pd.DataFrame([r.to_dict() for r in records])
    session.close()

    if len(df) == 0:
        st.markdown("""
        <div class="card" style="padding: 40px; text-align: center;">
            <div style="font-size: 48px; margin-bottom: 16px;">📋</div>
            <div style="font-size: 18px; color: #333333; margin-bottom: 8px;">暂无数据</div>
            <div style="font-size: 14px; color: #666666;">请前往「数据导入」上传数据，或继续「手动标注」已有数据</div>
        </div>
        """, unsafe_allow_html=True)
    else:
        total = len(df)
        valid = df['is_valid'].sum()
        s1_mask = df['comment_time'].astype(str).str.startswith('2026-04')
        s2_mask = df['comment_time'].astype(str).str.startswith('2026-05')
        labeled = df['cognitive_s2'].notna() & (df['cognitive_s2'] != '')

        cols = st.columns(4)
        metrics = [
            ("总数据量", total, "📦"),
            ("有效数据", valid, "✅"),
            ("无效数据", total - valid, "❌"),
            ("已标注", labeled.sum(), "🏷️")
        ]
        for col, (label, value, icon) in zip(cols, metrics):
            with col:
                render_metric_card(label, value, icon=icon)

        cols2 = st.columns(3)
        metrics2 = [
            ("第一阶段(4月)", int(s1_mask.sum()), "🗓️"),
            ("第二阶段(5月)", int(s2_mask.sum()), "🗓️"),
            ("有效率", f"{valid/total*100:.1f}%", "📝")
        ]
        for col, (label, value, icon) in zip(cols2, metrics2):
            with col:
                render_metric_card(label, value, icon=icon)

        st.markdown("<hr>", unsafe_allow_html=True)

        render_section_title("🧹 清洗规则统计", "🧹")

        rule_stats = []
        for code, name in [('pure_at', '纯@无互动'), ('pure_emoji', '纯表情'), ('filler_word', '无意义语气词'), ('empty', '空内容'), ('garbled', '乱码')]:
            col_name = f'is_{code}'
            cnt = df[col_name].sum() if col_name in df.columns else 0
            pct = cnt / total * 100
            rule_stats.append({"规则": name, "数量": int(cnt), "占比": pct})

        rule_df = pd.DataFrame(rule_stats)

        cols = st.columns(5)
        for i, (_, row) in enumerate(rule_df.iterrows()):
            with cols[i]:
                st.markdown(f"""
                <div class="card" style="padding: 16px; text-align: center;">
                    <div style="font-size: 11px; color: #666666; text-transform: uppercase;">{row['规则']}</div>
                    <div style="font-size: 28px; color: #2E7DCD; font-weight: 700; margin: 8px 0;">{int(row['数量'])}</div>
                    <div style="font-size: 12px; color: #999999;">{row['占比']:.1f}%</div>
                </div>
                """, unsafe_allow_html=True)

        st.markdown("<hr>", unsafe_allow_html=True)
        render_section_title("📋 数据列表", "📋")

        col1, col2, col3 = st.columns(3)
        with col1:
            show_filter = st.selectbox("", ["全部", "有效数据", "无效数据"], label_visibility="collapsed")
        with col2:
            stage_filter = st.selectbox("", ["全部", "阶段一(4月)", "阶段二(5月)"], label_visibility="collapsed")
        with col3:
            search = st.text_input("", placeholder="🔍 搜索内容...", label_visibility="collapsed")

        display_df = df.copy()
        if show_filter == "有效数据":
            display_df = display_df[display_df['is_valid'] == True]
        elif show_filter == "无效数据":
            display_df = display_df[display_df['is_valid'] == False]
        if stage_filter == "阶段一(4月)":
            display_df = display_df[s1_mask]
        elif stage_filter == "阶段二(5月)":
            display_df = display_df[s2_mask]
        if search:
            display_df = display_df[display_df['comment_content'].astype(str).str.contains(search, na=False)]

        st.markdown(f"<div style='color: #666666; margin: 16px 0;'>共 <span style='color: #2E7DCD;'>{len(display_df)}</span> 条数据</div>", unsafe_allow_html=True)

        page_size = 50
        total_pages = max(1, (len(display_df) - 1) // page_size + 1)
        page = st.number_input("", min_value=1, max_value=total_pages, value=1, label_visibility="collapsed")
        start_idx = (page - 1) * page_size
        page_df = display_df.iloc[start_idx:start_idx + page_size].copy()

        page_df['comment_content'] = page_df['comment_content'].astype(str).str[:60] + "..."
        page_df['is_valid'] = page_df['is_valid'].apply(lambda x: "✅" if x else "❌")

        st.dataframe(page_df[['id', 'comment_content', 'comment_time', 'is_valid', 'cognitive_s1', 'emotional_s1', 'action_s1']], hide_index=True, use_container_width=True, height=400)

# ============== 数据清洗页 ==============
elif page == "数据清洗":
    render_page_header("🧹 数据清洗", "配置清洗规则，实时预览效果")
    st.markdown("<p style='color: #666666;'>勾选清洗规则，实时预览效果</p>", unsafe_allow_html=True)

    session = get_session()
    records = session.query(RawData).filter(RawData.project_id == project_id).all()
    df = pd.DataFrame([r.to_dict() for r in records])
    rules = session.query(CleanRule).order_by(CleanRule.priority.desc()).all()
    session.close()

    if len(df) == 0:
        st.warning("⚠️ 请先导入数据")
        st.stop()

    render_section_title("🧹 清洗规则配置", "🧹")

    rule_enabled = {}
    cols = st.columns(3)
    for i, rule in enumerate(rules):
        with cols[i % 3]:
            enabled = st.checkbox(f"☑️ {rule.name}", value=rule.enabled, key=f"rule_{rule.code}")
            rule_enabled[rule.code] = enabled

    def calc_valid(row):
        if rule_enabled.get('pure_at', True) and row.get('is_pure_at'):
            return False
        if rule_enabled.get('pure_emoji', True) and row.get('is_pure_emoji'):
            return False
        if rule_enabled.get('filler_word', True) and row.get('is_filler_word'):
            return False
        if rule_enabled.get('empty', True) and row.get('is_empty'):
            return False
        if rule_enabled.get('garbled', True) and row.get('is_garbled'):
            return False
        return True

    df['_temp_valid'] = df.apply(calc_valid, axis=1)
    new_valid = df['_temp_valid'].sum()
    new_invalid = len(df) - new_valid

    st.markdown("<hr>", unsafe_allow_html=True)
    render_section_title("📊 清洗效果预览", "📊")

    cols = st.columns(4)
    metrics = [
        ("总数据量", len(df), "📦"),
        ("有效（启用后）", int(new_valid), "✅"),
        ("无效（剔除）", int(new_invalid), "❌"),
        ("有效率", f"{new_valid/len(df)*100:.1f}%", "📝")
    ]
    for col, (label, value, icon) in zip(cols, metrics):
        with col:
            render_metric_card(label, value, icon=icon)

    st.markdown("<hr>", unsafe_allow_html=True)

    cols = st.columns(2)
    with cols[0]:
        if st.button("💾 保存规则配置", use_container_width=True):
            session = get_session()
            for code, enabled in rule_enabled.items():
                rule = session.query(CleanRule).filter(CleanRule.code == code).first()
                if rule:
                    rule.enabled = enabled
            session.commit()
            session.close()
            st.success("规则配置已保存！")
    with cols[1]:
        if st.button("🔄 重新检测", use_container_width=True):
            with st.spinner("重新检测中..."):
                session = get_session()
                records = session.query(RawData).filter(RawData.project_id == project_id).all()
                for rd in records:
                    rd.is_pure_at = is_pure_at(rd.comment_content)
                    rd.is_pure_emoji = is_pure_emoji(rd.comment_content)
                    rd.is_filler_word = is_filler_word(rd.comment_content)
                    rd.is_empty = is_empty_content(rd.comment_content)
                    rd.is_garbled = is_garbled(rd.comment_content)
                    rd.is_valid = not any([rd.is_pure_at, rd.is_pure_emoji, rd.is_filler_word, rd.is_empty, rd.is_garbled])
                session.commit()
                session.close()
            st.success("重新检测完成！")
            st.rerun()

# ============== 手动标注页 ==============
elif page == "自动打标":
    render_page_header("⚙️ 自动打标", "根据当前规则批量生成认知、情绪、行动和品牌标签")

    project, df = get_project_dataframe(project_id)
    if len(df) == 0:
        st.warning("请先导入数据。")
        st.stop()

    session = get_session()
    rules = session.query(LabelRule).filter(LabelRule.enabled == True).all()
    session.close()
    by_layer = {}
    for rule in rules:
        by_layer[rule.layer] = by_layer.get(rule.layer, 0) + 1

    cols = st.columns(4)
    for col, item in zip(cols, [
        ("认知规则", by_layer.get("cognitive", 0), "🧠"),
        ("情绪规则", by_layer.get("emotional", 0), "😀"),
        ("行动规则", by_layer.get("action", 0), "🎯"),
        ("品牌规则", by_layer.get("brand", 0), "🏢"),
    ]):
        with col:
            render_metric_card(item[0], item[1], icon=item[2])

    valid_df = df[df['is_valid'] == True].copy()
    empty_label_count = int(((valid_df['cognitive_s1'].fillna("") == "") & (valid_df['cognitive_s2'].fillna("") == "")).sum())
    st.markdown("<div class='report-card'>", unsafe_allow_html=True)
    st.markdown("### 打标策略")
    st.markdown(f"<p class='quiet-note'>当前有效数据 {len(valid_df):,} 条，其中完全未打认知标签的样本约 {empty_label_count:,} 条。建议日常使用「只填空标签」，当规则大幅升级时再使用「覆盖重打」。</p>", unsafe_allow_html=True)
    action_cols = st.columns(2)
    with action_cols[0]:
        if st.button("只填空标签", type="primary", use_container_width=True):
            updated, total_valid = apply_auto_labels(project_id, overwrite=False)
            st.success(f"自动打标完成：有效数据 {total_valid:,} 条，填充字段 {updated:,} 个。")
            st.rerun()
    with action_cols[1]:
        if st.button("覆盖重打全部有效数据", use_container_width=True):
            updated, total_valid = apply_auto_labels(project_id, overwrite=True)
            st.success(f"覆盖重打完成：有效数据 {total_valid:,} 条，更新字段 {updated:,} 个。")
            st.rerun()
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div class='report-card'>", unsafe_allow_html=True)
    st.markdown("### 自动打标预览")
    preview_cols = ['comment_content', 'comment_time', 'cognitive_s1', 'cognitive_s2', 'emotional_s1', 'emotional_s2', 'action_s1', 'action_s2', 'brand_detected']
    st.dataframe(valid_df[preview_cols].head(100), hide_index=True, use_container_width=True, height=620)
    st.markdown("</div>", unsafe_allow_html=True)

# ============== 手动标注页 ==============
elif page == "手动标注":
    render_page_header("✏️ 手动标注", "在大表格中直接编辑标注，类似Excel操作体验")

    session = get_session()
    records = session.query(RawData).filter(RawData.project_id == project_id).all()
    df = pd.DataFrame([r.to_dict() for r in records])
    session.close()

    if len(df) == 0:
        st.warning("⚠️ 请先导入数据")
        st.stop()

    df_valid = df[df['is_valid'] == True].copy()

    # 筛选条件
    col1, col2, col3 = st.columns(3)
    with col1:
        stage_filter = st.selectbox("阶段", ["全部", "阶段一(4月)", "阶段二(5月)"], key="stage_filter")
    with col2:
        label_filter = st.selectbox("标注状态", ["全部", "未标注", "已标注"], key="label_filter")
    with col3:
        content_filter = st.selectbox("内容类型", ["全部", "普通内容", "长内容(>100字)", "提及竞品", "视频内容"], key="content_filter")

    filtered = df_valid.copy()
    if stage_filter == "阶段一(4月)":
        filtered = filtered[filtered['comment_time'].astype(str).str.startswith('2026-04')]
    elif stage_filter == "阶段二(5月)":
        filtered = filtered[filtered['comment_time'].astype(str).str.startswith('2026-05')]
    if label_filter == "未标注":
        filtered = filtered[filtered['cognitive_s2'].isna() | (filtered['cognitive_s2'] == '')]
    elif label_filter == "已标注":
        filtered = filtered[filtered['cognitive_s2'].notna() & (filtered['cognitive_s2'] != '')]
    if content_filter != "全部":
        filtered = filtered[filtered['content_type'] == content_filter]

    st.markdown(f"<div style='color: #666666; margin: 16px 0;'>共 <span style='color: #2E7DCD; font-weight: 600;'>{len(filtered)}</span> 条数据待标注</div>", unsafe_allow_html=True)

    # 准备编辑表格数据
    edit_df = filtered[['id', 'comment_content', 'comment_time', 'ip_location', 'content_type',
                        'cognitive_s1', 'cognitive_s2', 'emotional_s1', 'emotional_s2',
                        'action_s1', 'action_s2', 'brand_detected']].copy()

    edit_df.columns = ['ID', '评论内容', '评论时间', 'IP属地', '内容类型',
                       '认知-S1', '认知-S2', '情绪-S1', '情绪-S2',
                       '行动-S1', '行动-S2', '品牌']

    # 使用data_editor实现类Excel编辑
    editable = st.data_editor(
        edit_df,
        use_container_width=True,
        height=700,
        hide_index=True,
        column_config={
            "ID": st.column_config.NumberColumn("ID", width="small", disabled=True),
            "评论内容": st.column_config.TextColumn("评论内容", width="large"),
            "评论时间": st.column_config.TextColumn("评论时间", width="medium", disabled=True),
            "IP属地": st.column_config.TextColumn("IP属地", width="small", disabled=True),
            "内容类型": st.column_config.TextColumn("内容类型", width="medium", disabled=True),
            "认知-S1": st.column_config.SelectboxColumn("认知-S1", width="medium",
                options=["", "无明确认知", "信息混淆", "精准认知", "泛化抵触"]),
            "认知-S2": st.column_config.SelectboxColumn("认知-S2", width="medium",
                options=["", "无明确认知", "信息混淆", "精准认知", "泛化抵触"]),
            "情绪-S1": st.column_config.SelectboxColumn("情绪-S1", width="medium",
                options=["", "正面", "中性", "恐慌焦虑", "庆幸旁观", "愤怒背叛"]),
            "情绪-S2": st.column_config.SelectboxColumn("情绪-S2", width="medium",
                options=["", "正面", "中性", "恐慌焦虑", "庆幸旁观", "愤怒背叛"]),
            "行动-S1": st.column_config.SelectboxColumn("行动-S1", width="medium",
                options=["", "暂无行动", "寻求帮助", "转奶流失", "维权诉求"]),
            "行动-S2": st.column_config.SelectboxColumn("行动-S2", width="medium",
                options=["", "暂无行动", "寻求帮助", "转奶流失", "维权诉求"]),
            "品牌": st.column_config.TextColumn("品牌", width="medium"),
        }
    )

    # 保存按钮
    if st.button("💾 批量保存所有修改", type="primary", use_container_width=True):
        save_count = 0
        session = get_session()
        for idx, row in editable.iterrows():
            rd = session.query(RawData).filter(RawData.id == row['ID']).first()
            if rd:
                rd.cognitive_s1 = row['认知-S1'] if pd.notna(row['认知-S1']) else ''
                rd.cognitive_s2 = row['认知-S2'] if pd.notna(row['认知-S2']) else ''
                rd.emotional_s1 = row['情绪-S1'] if pd.notna(row['情绪-S1']) else ''
                rd.emotional_s2 = row['情绪-S2'] if pd.notna(row['情绪-S2']) else ''
                rd.action_s1 = row['行动-S1'] if pd.notna(row['行动-S1']) else ''
                rd.action_s2 = row['行动-S2'] if pd.notna(row['行动-S2']) else ''
                rd.brand_detected = row['品牌'] if pd.notna(row['品牌']) else ''
                rd.manual_override = True
                rd.labeled_by = st.session_state.get('user_id')
                rd.labeled_at = datetime.now()
                save_count += 1
        session.commit()
        session.close()
        st.success(f"✅ 已保存 {save_count} 条标注记录！")
        st.rerun()

    # 导出按钮 - 导出与参考文件相同格式的Excel
    if st.button("📥 导出Excel（参考格式）", use_container_width=True):
        # 读取原始数据并合并标注结果
        session = get_session()
        all_records = session.query(RawData).filter(RawData.project_id == project_id).all()
        session.close()

        # 构建导出数据
        export_data = []
        for rd in all_records:
            export_data.append({
                '视频ID': rd.video_id or '',
                '视频链接': rd.video_link or '',
                '关键词': rd.keyword or '',
                '内容': rd.content or '',
                '话题标签': rd.topic_tags or '',
                '发布时间': rd.publish_time or '',
                '博主名': rd.blogger or '',
                '收藏量': rd.favorites or 0,
                '评论量': rd.comments_count or 0,
                '点赞量': rd.likes or 0,
                '分享量': rd.shares or 0,
                '评论id': rd.comment_id or '',
                '评论时间': rd.comment_time or '',
                'ip_location': rd.ip_location or '',
                '评论内容': rd.comment_content or '',
                '昵称': rd.nickname or '',
                '二级评论数': rd.reply_count or 0,
                '评论获赞': rd.reply_likes or 0,
                '评论内容类型': rd.content_type or '',
                '认知层阶段一': rd.cognitive_s1 or '',
                '认知层阶段二': rd.cognitive_s2 or '',
                '情绪层阶段一': rd.emotional_s1 or '',
                '情绪层阶段二': rd.emotional_s2 or '',
                '行动层阶段一': rd.action_s1 or '',
                '行动层阶段二': rd.action_s2 or '',
                '品牌提及': rd.brand_detected or ''
            })

        export_df = pd.DataFrame(export_data)

        # 保存为Excel
        export_path = os.path.join(os.path.dirname(__file__), 'exports', f'标注导出_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx')
        os.makedirs(os.path.dirname(export_path), exist_ok=True)
        export_df.to_excel(export_path, index=False, engine='openpyxl')

        st.success(f"✅ 已导出 {len(export_df)} 条数据到 Excel！")
        with open(export_path, 'rb') as f:
            st.download_button(
                "📥 下载导出的Excel文件",
                data=f.read(),
                file_name=os.path.basename(export_path),
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )

    # CSV导出按钮
    st.download_button(
        "📥 导出CSV",
        data=editable.to_csv(index=False).encode('utf-8-sig'),
        file_name="舆情标注数据.csv",
        mime="text/csv",
        use_container_width=True
    )

# ============== 品牌竞品页 ==============
elif page == "品牌竞品":
    render_page_header("🏭 品牌与竞品管理", "管理品牌和竞品信息")

    tab1, tab2 = st.tabs(["🏢 品牌管理", "📦 竞品产品"])

    with tab1:
        session = get_session()
        brands = session.query(Brand).all()
        session.close()

        cols = st.columns(2)
        with cols[0]:
            st.markdown(f"""
            <div class="card" style="padding: 24px; text-align: center;">
                <div style="font-size: 12px; color: #666666; text-transform: uppercase;">品牌总数</div>
                <div style="font-size: 48px; color: #2E7DCD; font-weight: 700; margin-top: 8px;">{len(brands)}</div>
            </div>
            """, unsafe_allow_html=True)
        with cols[1]:
            active = sum(1 for b in brands if b.is_active)
            st.markdown(f"""
            <div class="card" style="padding: 24px; text-align: center;">
                <div style="font-size: 12px; color: #666666; text-transform: uppercase;">活跃品牌</div>
                <div style="font-size: 48px; color: #4CAF50; font-weight: 700; margin-top: 8px;">{active}</div>
            </div>
            """, unsafe_allow_html=True)

        st.markdown("<hr>", unsafe_allow_html=True)

        with st.expander("➕ 添加新品牌"):
            col1, col2 = st.columns(2)
            with col1:
                new_name = st.text_input("品牌名称")
                new_alias = st.text_input("别名（用|分隔）")
            with col2:
                new_category = st.selectbox("类别", ["奶粉", "辅食", "营养品", "纸尿裤", "其他"])
                new_priority = st.slider("优先级", 0, 10, 5)

            if st.button("添加品牌"):
                if new_name:
                    session = get_session()
                    brand = Brand(name=new_name, alias=new_alias, category=new_category, priority=new_priority)
                    session.add(brand)
                    session.commit()
                    session.close()
                    st.success(f"已添加品牌「{new_name}」")
                    st.rerun()

        brand_data = []
        for b in brands:
            brand_data.append({
                "ID": b.id,
                "品牌名": b.name,
                "别名": b.alias,
                "类别": b.category,
                "优先级": b.priority,
                "状态": "✅ 活跃" if b.is_active else "❌ 禁用"
            })
        st.dataframe(pd.DataFrame(brand_data), hide_index=True, use_container_width=True)

    with tab2:
        session = get_session()
        brands = session.query(Brand).filter(Brand.is_active == True).all()
        competitors = session.query(CompetitorProduct).all()
        session.close()

        cols = st.columns(2)
        with cols[0]:
            st.markdown(f"""
            <div class="card" style="padding: 24px; text-align: center;">
                <div style="font-size: 12px; color: #666666; text-transform: uppercase;">竞品总数</div>
                <div style="font-size: 48px; color: #2E7DCD; font-weight: 700; margin-top: 8px;">{len(competitors)}</div>
            </div>
            """, unsafe_allow_html=True)
        with cols[1]:
            active = sum(1 for c in competitors if c.is_active)
            st.markdown(f"""
            <div class="card" style="padding: 24px; text-align: center;">
                <div style="font-size: 12px; color: #666666; text-transform: uppercase;">活跃竞品</div>
                <div style="font-size: 48px; color: #4CAF50; font-weight: 700; margin-top: 8px;">{active}</div>
            </div>
            """, unsafe_allow_html=True)

        st.markdown("<hr>", unsafe_allow_html=True)

        with st.expander("➕ 添加竞品"):
            col1, col2 = st.columns(2)
            with col1:
                brand_choice = st.selectbox("所属品牌", [b.name for b in brands])
                prod_name = st.text_input("产品名称")
            with col2:
                prod_line = st.text_input("产品系列")
                prod_keywords = st.text_area("识别关键词（逗号分隔）")

            if st.button("添加竞品"):
                if prod_name and brand_choice:
                    session = get_session()
                    brand_obj = session.query(Brand).filter(Brand.name == brand_choice).first()
                    if brand_obj:
                        comp = CompetitorProduct(
                            brand_id=brand_obj.id,
                            name=prod_name,
                            product_line=prod_line,
                            keywords=[k.strip() for k in prod_keywords.split(',') if k.strip()]
                        )
                        session.add(comp)
                        session.commit()
                    session.close()
                    st.success(f"已添加竞品「{prod_name}」")
                    st.rerun()

        comp_data = []
        for c in competitors:
            brand_name = session.query(Brand).filter(Brand.id == c.brand_id).first().name if session.query(Brand).filter(Brand.id == c.brand_id).first() else "未知"
            comp_data.append({
                "ID": c.id,
                "产品名": c.name,
                "品牌": brand_name,
                "系列": c.product_line,
                "关键词": ", ".join(c.keywords or []),
                "状态": "✅ 活跃" if c.is_active else "❌ 禁用"
            })
        st.dataframe(pd.DataFrame(comp_data), hide_index=True, use_container_width=True)

# ============== 规则学习页 ==============
elif page == "规则学习":
    render_page_header("🧠 规则学习", "从人工校正样本中提取可确认的规则优化建议")

    project, df = get_project_dataframe(project_id)
    if len(df) == 0:
        st.warning("请先导入并完成部分人工校正。")
        st.stop()

    manual_count = int(df['manual_override'].fillna(False).sum())
    st.markdown(f"""
    <div class="report-card">
        <h3>学习逻辑</h3>
        <p class="quiet-note">系统会扫描人工校正过的样本，按「层级-阶段-标签」聚合评论中的候选词。如果候选词不在当前规则关键词中，且重复出现达到阈值，就作为建议展示。你确认后才会写入规则。</p>
        <span class="status-pill">人工校正样本 {manual_count:,} 条</span>
    </div>
    """, unsafe_allow_html=True)

    min_count = st.slider("建议关键词最小出现次数", min_value=1, max_value=10, value=2)
    suggestions = extract_learning_suggestions(project_id, min_count=min_count)

    if len(suggestions) == 0:
        st.info("暂时没有达到阈值的新关键词建议。可以先在「手动标注」里校正更多样本，或把阈值调低。")
    else:
        st.markdown("<div class='report-card'>", unsafe_allow_html=True)
        st.markdown("### 候选规则建议")
        edited_suggestions = st.data_editor(
            suggestions,
            hide_index=True,
            use_container_width=True,
            height=620,
            column_config={
                "层级": st.column_config.SelectboxColumn("层级", options=["cognitive", "emotional", "action"]),
                "阶段": st.column_config.SelectboxColumn("阶段", options=["s1", "s2"]),
                "标签": st.column_config.TextColumn("标签"),
                "建议关键词": st.column_config.TextColumn("建议关键词"),
                "出现次数": st.column_config.NumberColumn("出现次数", disabled=True),
            }
        )
        st.markdown("</div>", unsafe_allow_html=True)

        if st.button("✅ 将当前建议写入规则库", type="primary", use_container_width=True):
            session = get_session()
            added = 0
            for _, row in edited_suggestions.iterrows():
                layer = row["层级"]
                stage = row["阶段"]
                label = row["标签"]
                keyword = safe_text(row["建议关键词"]).strip()
                if not keyword:
                    continue
                rule = session.query(LabelRule).filter(
                    LabelRule.layer == layer,
                    LabelRule.stage == stage,
                    LabelRule.label == label
                ).order_by(LabelRule.priority.desc()).first()
                if rule:
                    keywords = rule.keywords or []
                    if keyword not in keywords:
                        keywords.append(keyword)
                        rule.keywords = keywords
                        added += 1
                else:
                    session.add(LabelRule(layer=layer, stage=stage, label=label, keywords=[keyword], priority=3, enabled=True))
                    added += 1
            session.commit()
            session.close()
            st.success(f"已写入 {added} 个新关键词。建议回到「自动打标」重新运行，并在「报告生成」确认结果。")
            st.rerun()

# ============== 其他页面（简化版） ==============
elif page == "标签管理":
    render_page_header("🏷️ 标签管理", "查看和管理标签统计")

    session = get_session()
    records = session.query(RawData).filter(RawData.project_id == project_id).all()
    df = pd.DataFrame([r.to_dict() for r in records])
    project = session.query(Project).filter(Project.id == project_id).first()
    session.close()

    if len(df) == 0:
        st.warning("⚠️ 请先导入数据")
        st.stop()

    df_valid = df[df['is_valid'] == True]
    s1_d = project.s1_denominator or 586
    s2_d = project.s2_denominator or 4093
    s1_mask = df_valid['comment_time'].astype(str).str.startswith('2026-04')
    s2_mask = df_valid['comment_time'].astype(str).str.startswith('2026-05')

    tab_cog, tab_emo, tab_act, tab_brand = st.tabs(["🧠 认知层", "😀 情绪层", "🎯 行动层", "🏢 品牌"])

    labels_map = {
        "认知层": ["无明确认知", "信息混淆", "精准认知", "泛化抵触"],
        "情绪层": ["中性", "正面", "恐慌焦虑", "庆幸旁观", "愤怒背叛"],
        "行动层": ["暂无行动", "寻求帮助", "转奶流失", "维权诉求"],
        "品牌": ["品牌提及", "竞品提及", "无品牌"]
    }

    for tab_name, labels, tab in [("认知层", labels_map["认知层"], tab_cog), ("情绪层", labels_map["情绪层"], tab_emo), ("行动层", labels_map["行动层"], tab_act), ("品牌", labels_map["品牌"], tab_brand)]:
        with tab:
            cols = st.columns(2)
            for i, (stage_name, stage_key, stage_mask, denom) in enumerate([("阶段一(4月)", "S1", s1_mask, s1_d), ("阶段二(5月)", "S2", s2_mask, s2_d)]):
                stage_df = df_valid[stage_mask]
                col_map = {"认知层": ('cognitive_s1', 'cognitive_s2'), "情绪层": ('emotional_s1', 'emotional_s2'), "行动层": ('action_s1', 'action_s2'), "品牌": ('brand_detected', 'brand_detected')}
                col = col_map.get(tab_name, ('cognitive_s1', 'cognitive_s2'))[0 if stage_key == 'S1' else 1]
                stats = stage_df[col].value_counts().to_dict() if col in stage_df.columns else {}

                with cols[i]:
                    st.markdown(f"""
                    <div class="card" style="padding: 20px;">
                        <div style="font-size: 14px; color: #2E7DCD; margin-bottom: 16px;">{stage_name} (总计: {denom} 条)</div>
                    """, unsafe_allow_html=True)
                    for label in labels:
                        count = stats.get(label, 0)
                        pct = count / denom * 100 if denom > 0 else 0
                        st.markdown(f"""
                        <div style="display: flex; justify-content: space-between; padding: 8px 0; border-bottom: 1px solid rgba(0,212,255,0.1);">
                            <span style="color: #444444;">{label}</span>
                            <span style="color: #2E7DCD;">{count} <span style="color: #999999;">({pct:.1f}%)</span></span>
                        </div>
                        """, unsafe_allow_html=True)
                    st.markdown("</div>", unsafe_allow_html=True)

elif page == "系统设置":
    render_page_header("⚙️ 系统设置", "系统配置和个人信息")

    st.markdown("""
    <div class="card" style="padding: 24px; margin: 16px 0;">
        <h3 style="color: #2E7DCD; margin-bottom: 16px;">👤 个人信息</h3>
        <p style="color: #666666;">用户名: {}</p>
        <p style="color: #666666;">角色: {}</p>
    </div>
    """.format(st.session_state.get('username', ''), st.session_state.get('role', '').upper()), unsafe_allow_html=True)

    st.markdown("""
    <div class="card" style="padding: 24px; margin: 16px 0;">
        <h3 style="color: #2E7DCD; margin-bottom: 16px;">🔧 系统信息</h3>
        <p style="color: #666666;">系统版本: 3.0.0</p>
        <p style="color: #666666;">数据库版本: SQLite</p>
    </div>
    """, unsafe_allow_html=True)

elif page == "用户管理" and st.session_state.get('role') == 'admin':
    render_page_header("👤 用户管理", "管理系统用户")

    session = get_session()
    users = session.query(User).all()
    session.close()

    user_data = []
    for u in users:
        user_data.append({
            "ID": u.id,
            "用户名": u.username,
            "显示名": u.display_name or "-",
            "角色": u.role,
            "状态": "✅ 活跃" if u.is_active else "❌ 禁用",
            "最后登录": u.last_login or "从未登录"
        })
    st.dataframe(pd.DataFrame(user_data), hide_index=True, use_container_width=True)

    with st.expander("➕ 添加用户"):
        col1, col2 = st.columns(2)
        with col1:
            new_username = st.text_input("用户名")
            new_password = st.text_input("密码", type="password")
        with col2:
            new_display = st.text_input("显示名称")
            new_role = st.selectbox("角色", ["user", "admin", "viewer"])

        if st.button("添加用户"):
            if new_username and new_password:
                session = get_session()
                user = User(username=new_username, password_hash=hash_password(new_password), display_name=new_display, role=new_role)
                session.add(user)
                session.commit()
                session.close()
                st.success("用户添加成功！")
                st.rerun()

# ============== 内容管理页 ==============
elif page == "内容管理":
    render_page_header("📋 内容管理", "多维度筛选内容，创建数据子集")
    st.markdown("<p style='color: #666666;'>多维度筛选内容，创建数据子集</p>", unsafe_allow_html=True)

    session = get_session()
    records = session.query(RawData).filter(RawData.project_id == project_id).all()
    df = pd.DataFrame([r.to_dict() for r in records])
    session.close()

    if len(df) == 0:
        st.warning("⚠️ 请先导入数据")
        st.stop()

    # 初始化筛选状态
    if 'content_filter_type' not in st.session_state:
        st.session_state['content_filter_type'] = 'all'

    render_section_title("🔍 筛选条件", "🔍")

    col1, col2, col3 = st.columns(3)
    with col1:
        stage = st.selectbox("阶段", ["全部", "阶段一(4月)", "阶段二(5月)"], key="content_stage")
    with col2:
        content_type = st.selectbox("内容类型", ["全部", "普通内容", "长内容(>100字)", "提及竞品", "视频内容"], key="content_type")
    with col3:
        label_status = st.selectbox("标注状态", ["全部", "已标注", "未标注"], key="content_label")

    col1, col2 = st.columns(2)
    with col1:
        keyword = st.text_input("🔍 关键词搜索", placeholder="输入关键词...", key="content_keyword")
    with col2:
        ip_loc = st.text_input("📍 IP地区", placeholder="如：北京、上海...", key="content_ip")

    # 应用筛选条件
    filtered = df.copy()
    if stage == "阶段一(4月)":
        filtered = filtered[filtered['comment_time'].astype(str).str.startswith('2026-04')]
    elif stage == "阶段二(5月)":
        filtered = filtered[filtered['comment_time'].astype(str).str.startswith('2026-05')]
    if content_type != "全部":
        filtered = filtered[filtered['content_type'] == content_type]
    if label_status == "已标注":
        filtered = filtered[filtered['cognitive_s2'].notna() & (filtered['cognitive_s2'] != '')]
    elif label_status == "未标注":
        filtered = filtered[(filtered['cognitive_s2'].isna()) | (filtered['cognitive_s2'] == '')]
    if keyword:
        filtered = filtered[filtered['comment_content'].astype(str).str.contains(keyword, na=False)]
    if ip_loc:
        filtered = filtered[filtered['ip_location'].astype(str).str.contains(ip_loc, na=False)]

    st.markdown("<hr>", unsafe_allow_html=True)

    # 可点击的筛选结果标签
    total_count = len(filtered)
    valid_count = int(filtered['is_valid'].sum())
    invalid_count = total_count - valid_count
    labeled_count = int((filtered['cognitive_s2'].notna() & (filtered['cognitive_s2'] != '')).sum())

    st.markdown("<div style='margin-bottom: 12px;'><span style='color: #666666; font-size: 14px;'>筛选结果：</span></div>", unsafe_allow_html=True)

    filter_cols = st.columns(4)
    with filter_cols[0]:
        if st.button(f"📦 总数据 ({total_count})", use_container_width=True):
            st.session_state['content_filter_type'] = 'all'
    with filter_cols[1]:
        if st.button(f"✅ 有效数据 ({valid_count})", use_container_width=True):
            filtered = filtered[filtered['is_valid'] == True]
            st.session_state['content_filter_type'] = 'valid'
    with filter_cols[2]:
        if st.button(f"❌ 无效数据 ({invalid_count})", use_container_width=True):
            filtered = filtered[filtered['is_valid'] == False]
            st.session_state['content_filter_type'] = 'invalid'
    with filter_cols[3]:
        if st.button(f"🏷️ 已标注 ({labeled_count})", use_container_width=True):
            filtered = filtered[filtered['cognitive_s2'].notna() & (filtered['cognitive_s2'] != '')]
            st.session_state['content_filter_type'] = 'labeled'

    st.markdown(f"<div style='color: #666666; margin: 16px 0;'>当前显示 <span style='color: #2E7DCD; font-weight: 600;'>{len(filtered)}</span> 条数据</div>", unsafe_allow_html=True)

    # 可编辑的数据表格 - 支持双击编辑
    edit_df = filtered[['id', 'comment_content', 'comment_time', 'ip_location', 'content_type',
                        'cognitive_s1', 'cognitive_s2', 'emotional_s1', 'emotional_s2',
                        'action_s1', 'action_s2']].copy()

    edit_df.columns = ['ID', '评论内容', '评论时间', 'IP属地', '内容类型',
                       '认知-S1', '认知-S2', '情绪-S1', '情绪-S2', '行动-S1', '行动-S2']

    edited_df = st.data_editor(
        edit_df,
        use_container_width=True,
        height=700,  # 约20行数据
        hide_index=True,
        column_config={
            "ID": st.column_config.NumberColumn("ID", width="small", disabled=True),
            "评论内容": st.column_config.TextColumn("评论内容", width="large"),
            "评论时间": st.column_config.TextColumn("评论时间", width="medium", disabled=True),
            "IP属地": st.column_config.TextColumn("IP属地", width="small", disabled=True),
            "内容类型": st.column_config.TextColumn("内容类型", width="medium", disabled=True),
            "认知-S1": st.column_config.SelectboxColumn("认知-S1", width="medium",
                options=["", "无明确认知", "信息混淆", "精准认知", "泛化抵触"]),
            "认知-S2": st.column_config.SelectboxColumn("认知-S2", width="medium",
                options=["", "无明确认知", "信息混淆", "精准认知", "泛化抵触"]),
            "情绪-S1": st.column_config.SelectboxColumn("情绪-S1", width="medium",
                options=["", "正面", "中性", "恐慌焦虑", "庆幸旁观", "愤怒背叛"]),
            "情绪-S2": st.column_config.SelectboxColumn("情绪-S2", width="medium",
                options=["", "正面", "中性", "恐慌焦虑", "庆幸旁观", "愤怒背叛"]),
            "行动-S1": st.column_config.SelectboxColumn("行动-S1", width="medium",
                options=["", "暂无行动", "寻求帮助", "转奶流失", "维权诉求"]),
            "行动-S2": st.column_config.SelectboxColumn("行动-S2", width="medium",
                options=["", "暂无行动", "寻求帮助", "转奶流失", "维权诉求"]),
        }
    )

    # 保存按钮
    if st.button("💾 批量保存所有修改", type="primary", use_container_width=True):
        save_count = 0
        session = get_session()
        for idx, row in edited_df.iterrows():
            rd = session.query(RawData).filter(RawData.id == row['ID']).first()
            if rd:
                rd.cognitive_s1 = row['认知-S1'] if pd.notna(row['认知-S1']) else ''
                rd.cognitive_s2 = row['认知-S2'] if pd.notna(row['认知-S2']) else ''
                rd.emotional_s1 = row['情绪-S1'] if pd.notna(row['情绪-S1']) else ''
                rd.emotional_s2 = row['情绪-S2'] if pd.notna(row['情绪-S2']) else ''
                rd.action_s1 = row['行动-S1'] if pd.notna(row['行动-S1']) else ''
                rd.action_s2 = row['行动-S2'] if pd.notna(row['行动-S2']) else ''
                rd.manual_override = True
                rd.labeled_by = st.session_state.get('user_id')
                rd.labeled_at = datetime.now()
                save_count += 1
        session.commit()
        session.close()
        st.success(f"✅ 已保存 {save_count} 条修改记录！")
        st.rerun()

    st.markdown("<hr>", unsafe_allow_html=True)
    render_section_title("📋 创建数据子集", "📋")

    col1, col2 = st.columns(2)
    with col1:
        subset_name = st.text_input("子集名称", placeholder="输入子集名称...")
    with col2:
        subset_desc = st.text_input("子集描述", placeholder="输入子集描述...")

    if st.button("📋 创建子集", use_container_width=True):
        if subset_name:
            session = get_session()
            subset = DataSubset(
                project_id=project_id,
                name=subset_name,
                description=subset_desc,
                data_ids=json.dumps(filtered['id'].tolist()),
                record_count=len(filtered)
            )
            session.add(subset)
            session.commit()
            session.close()
            st.success(f"✅ 子集「{subset_name}」创建成功！共 {len(filtered)} 条数据")
        else:
            st.warning("请输入子集名称")

# ============== 数据子集页 ==============
elif page == "数据子集":
    render_page_header("📋 数据子集管理", "管理已创建的数据子集")
    st.markdown("<p style='color: #666666;'>管理已创建的数据子集</p>", unsafe_allow_html=True)

    session = get_session()
    subsets = session.query(DataSubset).filter(DataSubset.project_id == project_id).all()
    session.close()

    if len(subsets) == 0:
        st.info("📭 暂无数据子集，请从「内容管理」创建")
    else:
        cols = st.columns(3)
        for i, sub in enumerate(subsets):
            with cols[i % 3]:
                st.markdown(f"""
                <div class="card" style="padding: 20px;">
                    <div style="font-size: 16px; color: #2E7DCD; font-weight: 600; margin-bottom: 8px;">{sub.name}</div>
                    <div style="font-size: 12px; color: #666666; margin-bottom: 12px;">{sub.description or '无描述'}</div>
                    <div style="display: flex; justify-content: space-between;">
                        <span style="color: #999999; font-size: 11px;">记录数</span>
                        <span style="color: #4CAF50; font-size: 14px; font-weight: 600;">{sub.record_count}</span>
                    </div>
                </div>
                """, unsafe_allow_html=True)

# ============== 规则管理页 ==============
elif page == "规则管理":
    render_page_header("📝 规则管理", "配置标签规则和清洗规则")

    tab1, tab2, tab3, tab4 = st.tabs(["🏷️ 标签规则", "🧹 清洗规则", "📊 优先级配置", "ℹ️ 使用说明"])

    with tab1:
        session = get_session()
        rules = session.query(LabelRule).all()
        session.close()

        # 层和阶段选择
        col_layer, col_stage = st.columns(2)
        with col_layer:
            layer = st.selectbox("选择层级", ["全部", "认知层", "情绪层", "行动层", "品牌竞品"])
        with col_stage:
            stage = st.selectbox("选择阶段", ["全部", "阶段一(S1)", "阶段二(S2)"])

        # 过滤规则
        stage_map = {"阶段一(S1)": "s1", "阶段二(S2)": "s2", "全部": None}
        stage_val = stage_map.get(stage)
        layer_map = {"认知层": "cognitive", "情绪层": "emotional", "行动层": "action", "品牌竞品": "brand"}
        filtered_rules = rules
        if layer != "全部":
            filtered_rules = [r for r in filtered_rules if r.layer == layer_map.get(layer, layer)]
        if stage_val:
            filtered_rules = [r for r in filtered_rules if r.stage == stage_val]

        st.markdown("""
        <div class="card" style="padding: 16px; margin: 12px 0;">
            <h3 style="color: #2E7DCD; margin-bottom: 12px;">标签规则说明</h3>
            <p style="color: #666666; line-height: 1.6; font-size: 13px;">
            <b style="color: #4CAF50;">认知层</b>：无明确认知 → 信息混淆 → 精准认知 → 泛化抵触（优先级依次升高）<br>
            <b style="color: #4CAF50;">情绪层</b>：中性 → 正面 → 恐慌焦虑 → 庆幸旁观 → 愤怒背叛<br>
            <b style="color: #4CAF50;">行动层</b>：暂无行动 → 寻求帮助 → 转奶流失 → 维权诉求
            </p>
        </div>
        """, unsafe_allow_html=True)

        if len(filtered_rules) == 0:
            st.info("暂无标签规则")
        else:
            # 按层级+阶段分组展示
            layer_stage_groups = {}
            for r in filtered_rules:
                key = f"{r.layer}_{r.stage}"
                if key not in layer_stage_groups:
                    layer_stage_groups[key] = []
                layer_stage_groups[key].append(r)

            for key, group_rules in layer_stage_groups.items():
                layer_name = {"cognitive": "认知层", "emotional": "情绪层", "action": "行动层", "brand": "品牌竞品"}.get(group_rules[0].layer, group_rules[0].layer)
                stage_name = "S1" if group_rules[0].stage == "s1" else "S2"
                with st.expander(f"📋 {layer_name} - {stage_name} ({len(group_rules)}条规则)", expanded=True):
                    # 显示每条规则的完整关键词
                    for r in sorted(group_rules, key=lambda x: x.priority, reverse=True):
                        keywords_str = ", ".join(r.keywords) if r.keywords else "(无关键词)"
                        st.markdown(f"""
                        <div style="background: #F5F7FA; border-radius: 6px; padding: 12px; margin: 8px 0; border-left: 3px solid #2E7DCD;">
                            <div style="display: flex; justify-content: space-between; align-items: center;">
                                <span style="font-weight: 600; color: #1A3A5C;">{r.label}</span>
                                <span style="font-size: 12px; color: #666;">优先级: {r.priority} | {'✅ 启用' if r.enabled else '❌ 禁用'}</span>
                            </div>
                            <div style="font-size: 12px; color: #666; margin-top: 6px;">关键词: {keywords_str}</div>
                        </div>
                        """, unsafe_allow_html=True)

            st.markdown("### ✍️ 规则编辑表格")
            rules_edit_df = pd.DataFrame([{
                "ID": r.id,
                "层级": r.layer,
                "阶段": r.stage,
                "标签": r.label,
                "关键词": ", ".join(r.keywords or []),
                "优先级": r.priority,
                "启用": bool(r.enabled),
            } for r in filtered_rules])
            edited_rules = st.data_editor(
                rules_edit_df,
                hide_index=True,
                use_container_width=True,
                height=520,
                column_config={
                    "ID": st.column_config.NumberColumn("ID", disabled=True),
                    "层级": st.column_config.SelectboxColumn("层级", options=["cognitive", "emotional", "action", "brand"]),
                    "阶段": st.column_config.SelectboxColumn("阶段", options=["s1", "s2"]),
                    "标签": st.column_config.TextColumn("标签"),
                    "关键词": st.column_config.TextColumn("关键词（逗号分隔）", width="large"),
                    "优先级": st.column_config.NumberColumn("优先级", min_value=0, max_value=100),
                    "启用": st.column_config.CheckboxColumn("启用"),
                },
                key="label_rules_editor"
            )
            if st.button("💾 保存规则表格修改", type="primary", use_container_width=True):
                session = get_session()
                for _, row in edited_rules.iterrows():
                    rule = session.query(LabelRule).filter(LabelRule.id == int(row["ID"])).first()
                    if not rule:
                        continue
                    rule.layer = row["层级"]
                    rule.stage = row["阶段"]
                    rule.label = row["标签"]
                    rule.keywords = [k.strip() for k in safe_text(row["关键词"]).split(",") if k.strip()]
                    rule.priority = to_int(row["优先级"], 0)
                    rule.enabled = bool(row["启用"])
                session.commit()
                session.close()
                st.success("规则修改已保存。")
                st.rerun()

        # 添加新规则
        st.markdown("---")
        st.markdown("### ➕ 添加新规则")
        with st.form("add_rule_form"):
            col1, col2, col3 = st.columns(3)
            with col1:
                new_layer = st.selectbox("层级", ["cognitive", "emotional", "action", "brand"], format_func=lambda x: {"cognitive": "认知层", "emotional": "情绪层", "action": "行动层", "brand": "品牌竞品"}.get(x, x))
            with col2:
                new_stage = st.selectbox("阶段", ["s1", "s2"], format_func=lambda x: "S1" if x == "s1" else "S2")
            with col3:
                new_priority = st.number_input("优先级", min_value=1, max_value=10, value=5)

            col_label, col_keywords = st.columns([1, 3])
            with col_label:
                new_label = st.text_input("标签名称")
            with col_keywords:
                new_keywords_str = st.text_input("关键词（逗号分隔）")

            col_submit, col_clear = st.columns([1, 4])
            with col_submit:
                submitted = st.form_submit_button("添加规则")
            if submitted:
                if new_label:
                    keywords_list = [k.strip() for k in new_keywords_str.split(",") if k.strip()] if new_keywords_str else []
                    session = get_session()
                    lr = LabelRule(
                        layer=new_layer,
                        stage=new_stage,
                        label=new_label,
                        keywords=keywords_list,
                        priority=new_priority,
                        enabled=True
                    )
                    session.add(lr)
                    session.commit()
                    session.close()
                    st.success(f"规则「{new_label}」添加成功！")
                    st.rerun()

        # 批量导入规则
        st.markdown("---")
        st.markdown("### 📥 批量导入规则（JSON格式）")
        with st.expander("批量导入"):
            json_input = st.text_area("粘贴JSON格式规则", height=150, placeholder='[{"layer": "emotional", "stage": "s1", "label": "恐慌焦虑", "keywords": ["吐奶", "拉肚子"], "priority": 3}]')
            if st.button("导入"):
                if json_input:
                    try:
                        rules_data = json.loads(json_input)
                        session = get_session()
                        for rule in rules_data:
                            lr = LabelRule(
                                layer=rule.get("layer"),
                                stage=rule.get("stage"),
                                label=rule.get("label"),
                                keywords=rule.get("keywords", []),
                                priority=rule.get("priority", 1),
                                enabled=rule.get("enabled", True)
                            )
                            session.add(lr)
                        session.commit()
                        session.close()
                        st.success(f"成功导入 {len(rules_data)} 条规则！")
                        st.rerun()
                    except json.JSONDecodeError:
                        st.error("JSON格式错误，请检查格式")

    with tab2:
        session = get_session()
        clean_rules = session.query(CleanRule).order_by(CleanRule.priority.desc()).all()
        session.close()

        render_section_title("🧹 清洗规则配置", "🧹")

        rules_data = []
        for r in clean_rules:
            rules_data.append({
                "ID": r.id,
                "规则名称": r.name,
                "规则代码": r.code,
                "描述": r.description or "-",
                "优先级": r.priority,
                "状态": "✅ 启用" if r.enabled else "❌ 禁用"
            })
        st.dataframe(pd.DataFrame(rules_data), hide_index=True, use_container_width=True)

        with st.expander("➕ 添加清洗规则"):
            col1, col2 = st.columns(2)
            with col1:
                cr_name = st.text_input("规则名称")
                cr_code = st.text_input("规则代码")
            with col2:
                cr_desc = st.text_input("规则描述")
                cr_priority = st.number_input("优先级", min_value=0, max_value=100, value=5)

            if st.button("添加规则"):
                if cr_name and cr_code:
                    session = get_session()
                    cr = CleanRule(name=cr_name, code=cr_code, description=cr_desc, priority=cr_priority)
                    session.add(cr)
                    session.commit()
                    session.close()
                    st.success("规则添加成功！")
                    st.rerun()

    with tab3:
        st.markdown("""
        <div class="card" style="padding: 24px; margin: 16px 0;">
            <h3 style="color: #2E7DCD; margin-bottom: 20px;">📊 标签匹配优先级</h3>

            <h4 style="color: #4CAF50; margin: 16px 0 8px;">认知层优先级</h4>
            <div style="color: #555555; line-height: 1.8;">
            1. <b>泛化抵触</b> > 其他（所有奶粉都有问题、企业欺骗等）<br>
            2. <b>精准认知</b> > 无明确认知（提到FDA、批次、国标等具体信息）<br>
            3. <b>信息混淆</b> > 无明确认知（分不清版本区别）<br>
            4. 默认：无明确认知
            </div>

            <h4 style="color: #4CAF50; margin: 20px 0 8px;">情绪层优先级</h4>
            <div style="color: #555555; line-height: 1.8;">
            1. <b>愤怒背叛</b> > 其他（企业不负责任、双标、欺骗）<br>
            2. <b>恐慌焦虑</b> > 中性（担忧宝宝健康）<br>
            3. <b>正面</b> > 中性（对a2持信任态度）<br>
            4. <b>庆幸旁观</b> > 中性（庆幸自己没买/已换）<br>
            5. 默认：中性
            </div>

            <h4 style="color: #4CAF50; margin: 20px 0 8px;">行动层优先级</h4>
            <div style="color: #555555; line-height: 1.8;">
            1. <b>维权诉求</b> > 其他（12315、投诉、赔偿）<br>
            2. <b>转奶流失</b> > 寻求帮助 > 暂无行动<br>
            3. <b>寻求帮助</b> > 暂无行动（询问怎么办、哪个安全）<br>
            4. 默认：暂无行动
            </div>
        </div>
        """, unsafe_allow_html=True)

    with tab4:
        st.markdown("""
        <div class="card" style="padding: 24px; margin: 16px 0;">
            <h3 style="color: #2E7DCD; margin-bottom: 20px;">ℹ️ 使用说明</h3>

            <h4 style="color: #4CAF50; margin: 16px 0 8px;">标签规则</h4>
            <p style="color: #666666; line-height: 1.8;">
            标签规则用于自动打标签，每个层级（认知/情绪/行动/品牌）都有对应的关键词配置。
            系统按优先级顺序匹配，匹配到即停止。
            </p>

            <h4 style="color: #4CAF50; margin: 20px 0 8px;">清洗规则</h4>
            <p style="color: #666666; line-height: 1.8;">
            清洗规则用于过滤无效数据，如纯@无互动、纯表情、乱码等。
            可以在数据清洗页面单独开关，实时预览效果。
            </p>

            <h4 style="color: #4CAF50; margin: 20px 0 8px;">分母配置</h4>
            <p style="color: #666666; line-height: 1.8;">
            每个项目的分母可配置，影响标签占比统计。
            第一阶段默认586条，第二阶段默认4093条（已剔除纯@无互动）。
            </p>
        </div>
        """, unsafe_allow_html=True)

# ============== 数据统计页 ==============
elif page == "数据统计":
    render_page_header("📊 数据统计", "查看标签分布和可视化图表")
    st.markdown("<p style='color: #666666;'>查看标签分布和可视化图表</p>", unsafe_allow_html=True)

    session = get_session()
    records = session.query(RawData).filter(RawData.project_id == project_id).all()
    df = pd.DataFrame([r.to_dict() for r in records])
    project = session.query(Project).filter(Project.id == project_id).first()
    session.close()

    if len(df) == 0:
        st.warning("⚠️ 请先导入数据")
        st.stop()

    df_valid = df[df['is_valid'] == True]
    s1_d = project.s1_denominator or 586
    s2_d = project.s2_denominator or 4093

    render_section_title("📊 统计配置", "📊")

    col1, col2, col3 = st.columns(3)
    with col1:
        stat_layer = st.selectbox("选择层级", ["认知层", "情绪层", "行动层"])
    with col2:
        stat_stage = st.selectbox("选择阶段", ["阶段一(4月)", "阶段二(5月)", "全部阶段"])
    with col3:
        chart_type = st.selectbox("图表类型", ["柱状图", "饼图"])

    layer_col_map = {"认知层": "cognitive_s2", "情绪层": "emotional_s2", "行动层": "action_s2"}
    denom_map = {"阶段一(4月)": s1_d, "阶段二(5月)": s2_d, "全部阶段": s1_d + s2_d}
    col = layer_col_map.get(stat_layer)
    denom = denom_map.get(stat_stage)

    if stat_stage == "全部阶段":
        s1_mask = df_valid['comment_time'].astype(str).str.startswith('2026-04')
        s2_mask = df_valid['comment_time'].astype(str).str.startswith('2026-05')
        s1_df = df_valid[s1_mask]
        s2_df = df_valid[s2_mask]
        s1_stats = s1_df[col].value_counts().to_dict() if col in s1_df.columns else {}
        s2_stats = s2_df[col].value_counts().to_dict() if col in s2_df.columns else {}

        labels_map = {"认知层": ["无明确认知", "信息混淆", "精准认知", "泛化抵触"],
                      "情绪层": ["中性", "正面", "恐慌焦虑", "庆幸旁观", "愤怒背叛"],
                      "行动层": ["暂无行动", "寻求帮助", "转奶流失", "维权诉求"]}

        cols = st.columns(2)
        for idx, (stage_name, stage_df, stage_stats, stage_denom) in enumerate([("阶段一(4月)", s1_df, s1_stats, s1_d), ("阶段二(5月)", s2_df, s2_stats, s2_d)]):
            with cols[idx]:
                st.markdown(f"""
                <div class="card" style="padding: 20px;">
                    <div style="font-size: 14px; color: #2E7DCD; margin-bottom: 16px;">{stage_name} (分母: {stage_denom})</div>
                """, unsafe_allow_html=True)
                for label in labels_map.get(stat_layer, []):
                    count = stage_stats.get(label, 0)
                    pct = count / stage_denom * 100 if stage_denom > 0 else 0
                    st.markdown(f"""
                    <div style="display: flex; justify-content: space-between; padding: 8px 0; border-bottom: 1px solid rgba(0,212,255,0.1);">
                        <span style="color: #444444;">{label}</span>
                        <span style="color: #2E7DCD;">{count} <span style="color: #999999;">({pct:.1f}%)</span></span>
                    </div>
                    """, unsafe_allow_html=True)
                st.markdown("</div>", unsafe_allow_html=True)
    else:
        stage_mask = df_valid['comment_time'].astype(str).str.startswith('2026-04' if '一' in stat_stage else '2026-05')
        stage_df = df_valid[stage_mask]
        stats = stage_df[col].value_counts().to_dict() if col in stage_df.columns else {}

        labels_map = {"认知层": ["无明确认知", "信息混淆", "精准认知", "泛化抵触"],
                      "情绪层": ["中性", "正面", "恐慌焦虑", "庆幸旁观", "愤怒背叛"],
                      "行动层": ["暂无行动", "寻求帮助", "转奶流失", "维权诉求"]}

        cols = st.columns(2)
        for idx, label in enumerate(labels_map.get(stat_layer, [])):
            with cols[idx % 2]:
                count = stats.get(label, 0)
                pct = count / denom * 100 if denom > 0 else 0
                st.markdown(f"""
                <div class="card" style="padding: 16px; margin: 8px 0;">
                    <div style="display: flex; justify-content: space-between; align-items: center;">
                        <span style="color: #444444; font-size: 14px;">{label}</span>
                        <span style="color: #2E7DCD; font-size: 20px; font-weight: 600;">{count}</span>
                    </div>
                    <div style="margin-top: 8px;">
                        <div class="progress-bar"><div class="fill" style="width: {pct}%;"></div></div>
                    </div>
                    <div style="text-align: right; font-size: 11px; color: #999999; margin-top: 4px;">{pct:.1f}% (分母: {denom})</div>
                </div>
                """, unsafe_allow_html=True)

    st.markdown("<hr>", unsafe_allow_html=True)
    render_section_title("📈 可视化图表", "📈")

    if stat_stage == "全部阶段":
        chart_df = pd.DataFrame([
            {"分类": label, "阶段一(4月)": s1_stats.get(label, 0), "阶段二(5月)": s2_stats.get(label, 0)}
            for label in labels_map.get(stat_layer, [])
        ])
    else:
        chart_df = pd.DataFrame([
            {"分类": label, "数量": stats.get(label, 0)}
            for label in labels_map.get(stat_layer, [])
        ])

    if chart_type == "柱状图":
        if stat_stage == "全部阶段":
            fig = px.bar(chart_df, x="分类", y=["阶段一(4月)", "阶段二(5月)"], barmode="group",
                        template="plotly_white", color_discrete_sequence=["#2E7DCD", "#4CAF50"])
        else:
            fig = px.bar(chart_df, x="分类", y="数量", template="plotly_white", color_discrete_sequence=["#2E7DCD"])
    else:
        fig = px.pie(chart_df, names="分类", values="数量" if stat_stage != "全部阶段" else "阶段一(4月)",
                    template="plotly_white", hole=0.4, color_discrete_sequence=["#2E7DCD", "#4CAF50", "#FF5252", "#FFAA00", "#9C27B0"])

    fig.update_layout(plot_bgcolor="#FFFFFF", paper_bgcolor="#FFFFFF",
                      font_color="#1A3A5C", margin=dict(l=20, r=20, t=40, b=20))
    st.plotly_chart(fig, use_container_width=True)

# ============== 报告生成页 ==============
elif page == "报告生成":
    render_page_header("📄 在线报告确认", "按最新版Word报告结构在线预览，确认后导出Word和参考格式Excel")

    project, df = get_project_dataframe(project_id)
    if len(df) == 0:
        st.warning("请先导入数据。")
        st.stop()

    context = build_report_context(project, df)
    confirmed_key = f"report_confirmed_{project_id}"

    cols = st.columns(5)
    for col, item in zip(cols, [
        ("总评论数", f"{context['total']:,}", "📦"),
        ("有效评论", f"{context['valid']:,}", "✅"),
        ("无效评论", f"{context['invalid']:,}", "🧹"),
        ("阶段一", f"{context['s1_count']:,}", "①"),
        ("阶段二", f"{context['s2_count']:,}", "②"),
    ]):
        with col:
            render_metric_card(item[0], item[1], icon=item[2])

    confirm_cols = st.columns([2, 1, 1])
    with confirm_cols[0]:
        st.markdown("<p class='quiet-note'>请先在下方在线报告中确认各板块数据和分析口径。确认后导出的Word会采用同一套统计结果。</p>", unsafe_allow_html=True)
    with confirm_cols[1]:
        if st.button("✅ 确认当前报告", type="primary", use_container_width=True):
            st.session_state[confirmed_key] = True
            st.success("已确认当前在线报告。")
    with confirm_cols[2]:
        if st.button("撤销确认", use_container_width=True):
            st.session_state[confirmed_key] = False
            st.info("已撤销确认，可继续修改标签和规则。")

    status = "已确认，可导出" if st.session_state.get(confirmed_key, False) else "待确认"
    st.markdown(f"<div class='report-card'><span class='status-pill'>报告状态：{status}</span></div>", unsafe_allow_html=True)

    st.markdown("<div class='report-card'>", unsafe_allow_html=True)
    st.markdown(f"### {context['project_name']} 舆情分析报告")
    st.markdown(f"<p class='quiet-note'>分析目的：帮助品牌了解消费者如何解读舆情事件，掌握用户情绪、认知变化、行动倾向和竞品提及情况。生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}</p>", unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div class='report-card'>", unsafe_allow_html=True)
    st.markdown("### 二、评论内容类型分布")
    if len(context["content_type"]) > 0:
        st.dataframe(context["content_type"], hide_index=True, use_container_width=True, height=260)
        fig = px.bar(context["content_type"], x="评论类型", y="数量", color_discrete_sequence=["#1A73E8"])
        fig.update_layout(template="plotly_white", margin=dict(l=20, r=20, t=20, b=20), height=320)
        st.plotly_chart(fig, use_container_width=True)
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("### 三、认知分析（Understanding）")
    c1, c2 = st.columns(2)
    with c1:
        render_stat_table("第一阶段：4月1日-5月1日 至初断货期", context["cognitive_s1"], context["s1_denominator"])
    with c2:
        render_stat_table("第二阶段：5月2日-5月18日 紫白金召回期", context["cognitive_s2"], context["s2_denominator"])

    st.markdown("### 四、情绪分析（Emotional）")
    c1, c2 = st.columns(2)
    with c1:
        render_stat_table("第一阶段：4月1日-5月1日 至初断货期", context["emotional_s1"], context["s1_denominator"])
    with c2:
        render_stat_table("第二阶段：5月2日-5月18日 紫白金召回期", context["emotional_s2"], context["s2_denominator"])

    st.markdown("### 五、行为分析（Action）")
    c1, c2 = st.columns(2)
    with c1:
        render_stat_table("第一阶段：4月1日-5月1日 至初断货期", context["action_s1"], context["s1_denominator"])
    with c2:
        render_stat_table("第二阶段：5月2日-5月18日 紫白金召回期", context["action_s2"], context["s2_denominator"])

    st.markdown("<div class='report-card'>", unsafe_allow_html=True)
    st.markdown("### 六、品牌竞品提及分析")
    if len(context["brand"]) > 0:
        st.dataframe(context["brand"], hide_index=True, use_container_width=True, height=360)
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div class='report-card'>", unsafe_allow_html=True)
    st.markdown("### 七、总结与建议")
    st.markdown("<p class='quiet-note'>建议优先复核恐慌焦虑、愤怒背叛、转奶流失和维权诉求样本；当人工校正量增加后，进入「规则学习」把高频修正线索写回规则库，再重新自动打标并确认报告。</p>", unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("---")
    export_cols = st.columns(2)
    filename_prefix = f"{context['project_name']}_{datetime.now().strftime('%Y%m%d')}"
    with export_cols[0]:
        doc_bytes = build_word_report(project, df, context)
        st.download_button(
            "📥 下载Word报告",
            data=doc_bytes,
            file_name=f"{filename_prefix}_舆情分析报告.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            disabled=not st.session_state.get(confirmed_key, False)
        )
    with export_cols[1]:
        session = get_session()
        records = session.query(RawData).filter(RawData.project_id == project_id).all()
        session.close()
        excel_bytes = dataframe_to_xlsx_bytes(build_export_dataframe(records))
        st.download_button(
            "📥 下载参考格式Excel",
            data=excel_bytes,
            file_name=f"{filename_prefix}_舆情数据.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True
        )

# ============== 模板管理页 ==============
elif page == "模板管理":
    render_page_header("📋 模板管理", "管理报告模板和板块配置")
    st.markdown("<p style='color: #666666;'>管理报告模板和板块配置</p>", unsafe_allow_html=True)

    session = get_session()
    templates = session.query(ReportTemplate).all()
    session.close()

    render_section_title("📋 模板列表", "📋")

    if len(templates) == 0:
        st.info("暂无报告模板")
    else:
        for t in templates:
            st.markdown(f"""
            <div class="card" style="padding: 20px; margin: 12px 0;">
                <div style="display: flex; justify-content: space-between; align-items: start;">
                    <div>
                        <div style="font-size: 16px; color: #2E7DCD; font-weight: 600;">{t.name}</div>
                        <div style="font-size: 12px; color: #999999; margin-top: 4px;">版本: {t.version} | 板块: {len(t.sections) if t.sections else 0}</div>
                        <div style="font-size: 12px; color: #666666; margin-top: 8px;">{t.description or '无描述'}</div>
                    </div>
                    <div class="tag {'tag-success' if t.is_default else 'tag-info'}">{'(默认)' if t.is_default else ''}</div>
                </div>
            </div>
            """, unsafe_allow_html=True)

    st.markdown("<hr>", unsafe_allow_html=True)
    render_section_title("➕ 创建新模板", "➕")

    col1, col2 = st.columns(2)
    with col1:
        tmpl_name = st.text_input("模板名称")
        tmpl_version = st.text_input("版本", value="1.0")
    with col2:
        tmpl_desc = st.text_input("模板描述")
        tmpl_default = st.checkbox("设为默认模板")

    st.markdown("**板块配置**")
    section_names = st.text_area("板块名称（每行一个）", placeholder="报告概述\n评论内容类型分布\n认知分析...")

    if st.button("💾 创建模板", use_container_width=True):
        if tmpl_name:
            session = get_session()
            sections = [{"id": f"sec_{i}", "name": name.strip(), "enabled": True, "order": i+1}
                       for i, name in enumerate(section_names.split('\n')) if name.strip()]
            tmpl = ReportTemplate(
                name=tmpl_name,
                version=tmpl_version,
                description=tmpl_desc,
                sections=sections,
                is_default=tmpl_default
            )
            session.add(tmpl)
            if tmpl_default:
                session.query(ReportTemplate).update({ReportTemplate.is_default: False})
            session.commit()
            session.close()
            st.success("模板创建成功！")
            st.rerun()
        else:
            st.warning("请输入模板名称")

# ============== 导出记录页 ==============
elif page == "导出记录":
    render_page_header("📁 导出记录", "查看所有导出历史记录")
    st.markdown("<p style='color: #666666;'>查看所有导出历史记录</p>", unsafe_allow_html=True)

    session = get_session()
    exports = session.query(ExportRecord).filter(ExportRecord.project_id == project_id).order_by(ExportRecord.exported_at.desc()).all()
    session.close()

    if len(exports) == 0:
        st.info("📭 暂无导出记录")
    else:
        render_section_title("📋 导出历史", "📋")

        export_data = []
        for e in exports:
            export_data.append({
                "ID": e.id,
                "文件名": e.filename,
                "模板": e.template_name or "-",
                "格式": e.export_type.upper(),
                "大小": f"{e.file_size/1024:.1f} KB" if e.file_size else "-",
                "时间": e.exported_at.strftime("%Y-%m-%d %H:%M") if e.exported_at else "-",
                "板块": len(e.included_sections) if e.included_sections else 0
            })
        st.dataframe(pd.DataFrame(export_data), hide_index=True, use_container_width=True)

        st.markdown("<hr>", unsafe_allow_html=True)
        render_section_title("⚙️ 文件名模板配置", "⚙️")

        st.markdown("""
        <div class="card" style="padding: 20px; margin: 16px 0;">
            <p style="color: #666666; line-height: 1.8;">
            <b style="color: #2E7DCD;">可用变量：</b><br>
            <code style="color: #4CAF50;">{project_name}</code> - 项目名称<br>
            <code style="color: #4CAF50;">{date}</code> - 生成日期<br>
            <code style="color: #4CAF50;">{version}</code> - 模板版本
            </p>
        </div>
        """, unsafe_allow_html=True)

        default_tpl = "{project_name}_{date}_{version}"
        new_tpl = st.text_input("文件名模板", value=default_tpl)

        if st.button("💾 保存模板", use_container_width=True):
            st.success("文件名模板已保存")

else:
    st.markdown(f"""
    <div style="text-align: center; padding: 100px 0;">
        <h2 style="color: rgba(255,255,255,0.3);">🚧 功能开发中</h2>
        <p style="color: rgba(0,0,0,0.08); margin-top: 16px;">页面: {page}</p>
    </div>
    """, unsafe_allow_html=True)
