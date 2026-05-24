"""
舆情分析 - 报告生成脚本
生成Word和Excel格式的分析报告
"""

import pandas as pd
from datetime import datetime
from pathlib import Path
from typing import Dict, List

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class ReportGenerator:
    """报告生成器"""

    def __init__(self, analysis_results: Dict, context: Dict = None):
        """
        初始化报告生成器

        Args:
            analysis_results: 分析结果字典
            context: 分析背景信息
        """
        self.results = analysis_results
        self.context = context or {}

    def generate_excel_report(self, output_file: str):
        """
        生成Excel格式报告

        Args:
            output_file: 输出文件路径
        """
        with pd.ExcelWriter(output_file, engine='openpyxl') as writer:
            # Sheet 1: 总览
            self._write_overview_sheet(writer)

            # Sheet 2: 情感分析详情
            self._write_sentiment_sheet(writer)

            # Sheet 3: 品牌提及分析
            self._write_brand_sheet(writer)

            # Sheet 4: 产品提及分析
            self._write_product_sheet(writer)

            # Sheet 5: 痛点词分析
            self._write_pain_points_sheet(writer)

        print(f"Excel报告已生成: {output_file}")

    def generate_word_report(self, output_file: str):
        """
        生成Word格式报告

        Args:
            output_file: 输出文件路径
        """
        if not HAS_DOCX:
            print("警告: python-docx未安装，无法生成Word报告")
            print("请运行: pip install python-docx")
            return

        doc = Document()

        # 标题
        title = doc.add_heading('舆情分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 分析时间
        doc.add_paragraph(f"分析时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"数据总量: {self.results.get('total_count', 0)}条")

        # 情感分析
        doc.add_heading('一、情感分析', level=1)
        sentiment_data = self.results.get('sentiment_distribution', {})
        if sentiment_data:
            table = doc.add_table(rows=len(sentiment_data) + 1, cols=3)
            table.style = 'Table Grid'
            # 表头
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '情感类型'
            hdr_cells[1].text = '数量'
            hdr_cells[2].text = '占比'
            # 数据
            for i, (sentiment, count) in enumerate(sentiment_data.items()):
                row_cells = table.rows[i + 1].cells
                row_cells[0].text = sentiment
                row_cells[1].text = str(count)
                row_cells[2].text = f"{count / sum(sentiment_data.values()) * 100:.1f}%"
        else:
            doc.add_paragraph('暂无数据')

        # 品牌提及分析
        doc.add_heading('二、品牌提及分析', level=1)
        brand_data = self.results.get('brand_mentions', {})
        if brand_data:
            table = doc.add_table(rows=len(brand_data) + 1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '品牌'
            hdr_cells[1].text = '提及次数'
            for i, (brand, count) in enumerate(brand_data.items()):
                row_cells = table.rows[i + 1].cells
                row_cells[0].text = brand
                row_cells[1].text = str(count)
        else:
            doc.add_paragraph('暂无数据')

        # 产品提及分析
        doc.add_heading('三、产品提及分析', level=1)
        product_data = self.results.get('product_mentions', {})
        if product_data:
            table = doc.add_table(rows=len(product_data) + 1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '产品'
            hdr_cells[1].text = '提及次数'
            for i, (product, count) in enumerate(product_data.items()):
                row_cells = table.rows[i + 1].cells
                row_cells[0].text = product
                row_cells[1].text = str(count)
        else:
            doc.add_paragraph('暂无数据')

        # 痛点词分析
        doc.add_heading('四、痛点词分析', level=1)
        pain_points = self.results.get('pain_points', {})
        if pain_points:
            for stage_kw, count in list(pain_points.items())[:20]:  # 限制显示前20条
                doc.add_paragraph(f"{stage_kw}: {count}次")
        else:
            doc.add_paragraph('暂无数据')

        # 分析结论
        doc.add_heading('五、分析结论', level=1)
        conclusion = self.generate_conclusion_text()
        doc.add_paragraph(conclusion)

        # 保存
        doc.save(output_file)
        print(f"Word报告已生成: {output_file}")

    def _write_overview_sheet(self, writer):
        """写入总览sheet"""
        data = {
            '指标': ['分析时间', '总数据量', '数据来源', '分析类型'],
            '值': [
                datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                self.results.get('total_count', 0),
                self.context.get('data_source', '待填写'),
                '舆情分析'
            ]
        }
        df = pd.DataFrame(data)
        df.to_excel(writer, sheet_name='总览', index=False)

    def _write_sentiment_sheet(self, writer):
        """写入情感分析sheet"""
        sentiment_dist = self.results.get('sentiment_distribution', {})
        if sentiment_dist:
            data = {
                '情感类型': list(sentiment_dist.keys()),
                '数量': list(sentiment_dist.values()),
                '占比': [f"{v/sum(sentiment_dist.values())*100:.1f}%" for v in sentiment_dist.values()]
            }
            df = pd.DataFrame(data)
            df.to_excel(writer, sheet_name='情感分析', index=False)

    def _write_brand_sheet(self, writer):
        """写入品牌提及sheet"""
        brand_mentions = self.results.get('brand_mentions', {})
        if brand_mentions:
            data = {
                '品牌': list(brand_mentions.keys()),
                '提及次数': list(brand_mentions.values()),
                '占比': [f"{v/sum(brand_mentions.values())*100:.1f}%" for v in brand_mentions.values()]
            }
            df = pd.DataFrame(data)
            df.to_excel(writer, sheet_name='品牌提及', index=False)

    def _write_product_sheet(self, writer):
        """写入产品提及sheet"""
        product_mentions = self.results.get('product_mentions', {})
        if product_mentions:
            data = {
                '产品': list(product_mentions.keys()),
                '提及次数': list(product_mentions.values()),
                '占比': [f"{v/sum(product_mentions.values())*100:.1f}%" for v in product_mentions.values()]
            }
            df = pd.DataFrame(data)
            df.to_excel(writer, sheet_name='产品提及', index=False)

    def _write_pain_points_sheet(self, writer):
        """写入痛点词分析sheet"""
        pain_points = self.results.get('pain_points', {})
        if pain_points:
            # 解析 stage:keyword 格式
            stages = {}
            for stage_kw, count in pain_points.items():
                if ':' in stage_kw:
                    stage, kw = stage_kw.split(':', 1)
                    if stage not in stages:
                        stages[stage] = []
                    stages[stage].append((kw, count))

            with pd.ExcelWriter(output_file, engine='openpyxl') as writer:
                for stage, items in stages.items():
                    safe_name = stage[:31]  # sheet名限制31字符
                    df = pd.DataFrame(items, columns=['关键词', '出现次数'])
                    df.to_excel(writer, sheet_name=safe_name, index=False)

    def generate_conclusion_text(self) -> str:
        """生成分析结论文本"""
        lines = []

        # 情感倾向总结
        sentiment = self.results.get('sentiment_distribution', {})
        if sentiment:
            total = sum(sentiment.values())
            positive_pct = sentiment.get('正面', 0) / total * 100 if total else 0
            negative_pct = sentiment.get('负面', 0) / total * 100 if total else 0

            if positive_pct > negative_pct:
                lines.append(f"整体情感倾向偏正面（正面{positive_pct:.1f}%），用户反馈良好。")
            elif negative_pct > positive_pct:
                lines.append(f"整体情感倾向偏负面（负面{negative_pct:.1f}%），需关注问题反馈。")
            else:
                lines.append("整体情感倾向较为中性，用户态度平稳。")

        # a2品牌提及
        a2_mentions = self.results.get('a2_mentions', 0)
        if a2_mentions > 0:
            lines.append(f"a2品牌被提及{a2_mentions}次。")

        # 痛点词总结
        pain_points = self.results.get('pain_points', {})
        if pain_points:
            top_pain = sorted(pain_points.items(), key=lambda x: x[1], reverse=True)[:5]
            top_str = '、'.join([f"{kw}({cnt}次)" for kw, cnt in top_pain])
            lines.append(f"高频痛点词：{top_str}。")

        return '\n'.join(lines) if lines else '数据量不足，暂无法生成结论。'

    def generate_conclusion(self) -> str:
        """生成分析结论（兼容旧接口）"""
        return self.generate_conclusion_text()


def generate_report(analysis_results: Dict, output_file: str, format: str = 'both', context: Dict = None):
    """
    生成报告主函数

    Args:
        analysis_results: 分析结果
        output_file: 输出文件路径（不含扩展名）
        format: 报告格式 ('excel', 'word', 'both')
        context: 分析背景信息
    """
    generator = ReportGenerator(analysis_results, context)

    if format.lower() in ['excel', 'both']:
        excel_file = output_file if output_file.endswith('.xlsx') else output_file + '.xlsx'
        generator.generate_excel_report(excel_file)
        print(f"Excel报告已生成: {excel_file}")

    if format.lower() in ['word', 'both']:
        word_file = output_file if output_file.endswith('.docx') else output_file + '.docx'
        generator.generate_word_report(word_file)
        print(f"Word报告已生成: {word_file}")

    # 输出结论
    print("\n=== 分析结论 ===")
    print(generator.generate_conclusion())


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 3:
        print("用法: python report.py <分析结果文件> <输出报告文件> [格式: excel|word|both]")
        print("默认格式: both（同时生成Excel和Word）")
    else:
        input_f = sys.argv[1]
        output_f = sys.argv[2]
        fmt = sys.argv[3] if len(sys.argv) > 3 else 'both'

        # 读取分析结果
        # TODO: 实现读取分析结果
        results = {}

        generate_report(results, output_f, fmt)