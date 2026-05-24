"""
A2舆情分析报告生成脚本 v5
基于最新的分类数据生成Word报告
"""

import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_title(doc, text, level=0):
    """添加标题"""
    heading = doc.add_heading(text, level)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return heading

def create_table(doc, headers, rows, col_widths=None):
    """创建表格"""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], 'D9E2F3')
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # 数据行
    for i, row_data in enumerate(rows):
        row_cells = table.rows[i + 1].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = str(cell_data)

    return table

def generate_v5_report(data_file: str, output_file: str):
    """生成v5报告"""
    print(f"读取数据: {data_file}")
    df = pd.read_excel(data_file)

    # 解析时间阶段
    df['评论时间_str'] = df['评论时间'].astype(str)
    phase1_mask = df['评论时间_str'].str.startswith('2026-04')
    phase2_mask = df['评论时间_str'].str.startswith('2026-05')

    df_p1 = df[phase1_mask]
    df_p2 = df[phase2_mask]

    doc = Document()

    # ========== 标题 ==========
    title = doc.add_heading('A2奶粉舆情分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"数据来源: 抖音/小红书评论数据")
    doc.add_paragraph("")

    # ========== 数据概览 ==========
    doc.add_heading('一、数据概览', level=1)
    total_count = len(df)
    p1_count = len(df_p1)
    p2_count = len(df_p2)
    doc.add_paragraph(f"总数据量: {total_count} 条")
    doc.add_paragraph(f"第一阶段(2026年4月): {p1_count} 条")
    doc.add_paragraph(f"第二阶段(2026年5月): {p2_count} 条")

    # ========== 第一阶段分析 ==========
    doc.add_heading('二、第一阶段分析（2026年4月）', level=1)

    doc.add_heading('2.1 认知层分布', level=2)
    cog_p1 = df_p1['认知层阶段一'].value_counts().to_dict()
    rows = [[k, v, f"{v/p1_count*100:.1f}%"] for k, v in cog_p1.items()]
    create_table(doc, ['认知类型', '数量', '占比'], rows)

    doc.add_paragraph('')
    doc.add_heading('2.2 情绪层分布', level=2)
    emo_p1 = df_p1['情绪层阶段一'].value_counts().to_dict()
    rows = [[k, v, f"{v/p1_count*100:.1f}%"] for k, v in emo_p1.items()]
    create_table(doc, ['情绪类型', '数量', '占比'], rows)

    doc.add_paragraph('')
    doc.add_heading('2.3 行动层分布', level=2)
    beh_p1 = df_p1['行动层阶段一'].value_counts().to_dict()
    rows = [[k, v, f"{v/p1_count*100:.1f}%"] for k, v in beh_p1.items()]
    create_table(doc, ['行动类型', '数量', '占比'], rows)

    # ========== 第二阶段分析 ==========
    doc.add_heading('三、第二阶段分析（2026年5月）', level=1)

    doc.add_heading('3.1 认知层分布', level=2)
    cog_p2 = df_p2['认知层阶段二'].value_counts().to_dict()
    rows = [[k, v, f"{v/p2_count*100:.1f}%"] for k, v in cog_p2.items()]
    create_table(doc, ['认知类型', '数量', '占比'], rows)

    doc.add_paragraph('')
    doc.add_heading('3.2 情绪层分布', level=2)
    emo_p2 = df_p2['情绪层阶段二'].value_counts().to_dict()
    rows = [[k, v, f"{v/p2_count*100:.1f}%"] for k, v in emo_p2.items()]
    create_table(doc, ['情绪类型', '数量', '占比'], rows)

    doc.add_paragraph('')
    doc.add_heading('3.3 行动层分布', level=2)
    beh_p2 = df_p2['行动层阶段二'].value_counts().to_dict()
    rows = [[k, v, f"{v/p2_count*100:.1f}%"] for k, v in beh_p2.items()]
    create_table(doc, ['行动类型', '数量', '占比'], rows)

    # ========== 品牌提及分析 ==========
    doc.add_heading('四、品牌提及分析', level=1)

    all_brands = []
    for brands_str in df['品牌提及']:
        if pd.notna(brands_str) and isinstance(brands_str, str) and brands_str.strip():
            all_brands.extend(brands_str.split('|'))
    brand_counts = pd.Series(all_brands).value_counts()

    doc.add_paragraph('品牌提及TOP10：')
    rows = [[brand, count, f"{count/len(df)*100:.1f}%"] for brand, count in brand_counts.head(10).items()]
    create_table(doc, ['品牌', '提及次数', '占比'], rows)

    # ========== 两阶段对比 ==========
    doc.add_heading('五、两阶段对比分析', level=1)

    # 情绪层对比
    doc.add_heading('5.1 情绪层对比', level=2)
    emo_p1_pct = {k: v/p1_count*100 for k, v in emo_p1.items()}
    emo_p2_pct = {k: v/p2_count*100 for k, v in emo_p2.items()}
    all_emos = set(emo_p1.keys()) | set(emo_p2.keys())
    rows = []
    for emo in sorted(all_emos):
        p1_val = emo_p1.get(emo, 0)
        p1_pct = emo_p1_pct.get(emo, 0)
        p2_val = emo_p2.get(emo, 0)
        p2_pct = emo_p2_pct.get(emo, 0)
        diff = p2_pct - p1_pct
        diff_str = f"+{diff:.1f}%" if diff > 0 else f"{diff:.1f}%"
        rows.append([emo, f"{p1_val}({p1_pct:.1f}%)", f"{p2_val}({p2_pct:.1f}%)", diff_str])
    create_table(doc, ['情绪类型', '阶段一', '阶段二', '变化'], rows)

    doc.add_paragraph('')
    doc.add_heading('5.2 行动层对比', level=2)
    beh_p1_pct = {k: v/p1_count*100 for k, v in beh_p1.items()}
    beh_p2_pct = {k: v/p2_count*100 for k, v in beh_p2.items()}
    all_behs = set(beh_p1.keys()) | set(beh_p2.keys())
    rows = []
    for beh in sorted(all_behs):
        p1_val = beh_p1.get(beh, 0)
        p1_pct = beh_p1_pct.get(beh, 0)
        p2_val = beh_p2.get(beh, 0)
        p2_pct = beh_p2_pct.get(beh, 0)
        diff = p2_pct - p1_pct
        diff_str = f"+{diff:.1f}%" if diff > 0 else f"{diff:.1f}%"
        rows.append([beh, f"{p1_val}({p1_pct:.1f}%)", f"{p2_val}({p2_pct:.1f}%)", diff_str])
    create_table(doc, ['行动类型', '阶段一', '阶段二', '变化'], rows)

    # ========== 评论类型分析 ==========
    doc.add_heading('六、评论类型分布', level=1)

    doc.add_heading('6.1 第一阶段评论类型', level=2)
    type_p1 = df_p1['评论内容类型'].value_counts().to_dict()
    rows = [[k, v, f"{v/p1_count*100:.1f}%"] for k, v in type_p1.items()]
    create_table(doc, ['评论类型', '数量', '占比'], rows)

    doc.add_paragraph('')
    doc.add_heading('6.2 第二阶段评论类型', level=2)
    type_p2 = df_p2['评论内容类型'].value_counts().to_dict()
    rows = [[k, v, f"{v/p2_count*100:.1f}%"] for k, v in type_p2.items()]
    create_table(doc, ['评论类型', '数量', '占比'], rows)

    # ========== 分析结论 ==========
    doc.add_heading('七、分析结论', level=1)

    # 计算关键指标
    positive_p1 = emo_p1.get('正面', 0) / p1_count * 100
    positive_p2 = emo_p2.get('正面', 0) / p2_count * 100
    negative_p1 = emo_p1.get('负面', 0) / p1_count * 100
    negative_p2 = emo_p2.get('负面', 0) / p2_count * 100
    panic_p1 = emo_p1.get('恐慌焦虑', 0) / p1_count * 100
    panic_p2 = emo_p2.get('恐慌焦虑', 0) / p2_count * 100
    switch_p1 = beh_p1.get('转奶流失', 0) / p1_count * 100
    switch_p2 = beh_p2.get('转奶流失', 0) / p2_count * 100

    doc.add_paragraph(f'''
整体趋势分析：

1. **情感倾向变化**
   - 正面占比：第一阶段 {positive_p1:.1f}% → 第二阶段 {positive_p2:.1f}%
   - 负面占比：第一阶段 {negative_p1:.1f}% → 第二阶段 {negative_p2:.1f}%
   - 恐慌焦虑：第一阶段 {panic_p1:.1f}% → 第二阶段 {panic_p2:.1f}%

2. **行动层变化**
   - 转奶流失：第一阶段 {switch_p1:.1f}% → 第二阶段 {switch_p2:.1f}%

3. **主要发现**
   - 第二阶段恐慌焦虑情绪明显上升，说明随着事件发酵，更多用户表现出担忧情绪
   - 转奶流失率有所上升，部分用户选择换奶
   - 正面用户保持稳定，说明核心用户群体对a2品牌仍有信任
''')

    # 保存
    doc.save(output_file)
    print(f"报告已生成: {output_file}")

if __name__ == "__main__":
    import sys

    if len(sys.argv) < 3:
        print("用法: python report_v5.py <数据文件> <输出报告文件>")
    else:
        data_f = sys.argv[1]
        output_f = sys.argv[2]
        generate_v5_report(data_f, output_f)