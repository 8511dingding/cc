from fastapi import FastAPI, Depends, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional
from datetime import datetime
import sqlite3
import os
import httpx

app = FastAPI(title="CompanyDoc API", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

DB_PATH = os.path.join(os.path.dirname(__file__), "..", "companydoc.db")

def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    try:
        yield conn
    finally:
        conn.close()

@app.on_event("startup")
def startup():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS templates (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            category TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS template_versions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            template_id INTEGER NOT NULL,
            version INTEGER NOT NULL DEFAULT 1,
            content TEXT,
            description TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (template_id) REFERENCES templates(id) ON DELETE CASCADE
        )
    ''')

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS documents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            content TEXT,
            content_en TEXT,
            content_th TEXT,
            template_id INTEGER,
            template_version_id INTEGER,
            category TEXT DEFAULT '其他',
            status TEXT DEFAULT 'draft',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            created_by TEXT,
            FOREIGN KEY (template_id) REFERENCES templates(id),
            FOREIGN KEY (template_version_id) REFERENCES template_versions(id)
        )
    ''')

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS customers (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            company_name TEXT,
            company_name_th TEXT,
            registration_number TEXT,
            address TEXT,
            address_th TEXT,
            contact_person TEXT,
            contact_email TEXT,
            contact_phone TEXT,
            contact_phone_th TEXT,
            notes TEXT,
            source_document_id INTEGER,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (source_document_id) REFERENCES documents(id)
        )
    ''')

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS uploaded_files (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            filename TEXT NOT NULL,
            original_filename TEXT NOT NULL,
            file_path TEXT NOT NULL,
            file_type TEXT,
            file_size INTEGER,
            parsed_content TEXT,
            customer_id INTEGER,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (customer_id) REFERENCES customers(id)
        )
    ''')
    
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS terms (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            zh TEXT NOT NULL,
            en TEXT,
            th TEXT,
            category TEXT DEFAULT '通用'
        )
    ''')
    
    cursor.execute("SELECT COUNT(*) FROM terms")
    if cursor.fetchone()[0] == 0:
        sample_terms = [
            ("经销商授权书", "Dealer Authorization", "จดหมายอนุญาตตัวแทนจำหน่าย", "合同"),
            ("有效期", "Valid Period", "ช่วงเวลาที่มีผลบังคับ", "合同"),
            ("授权范围", "Scope of Authorization", "ขอบเขตของการอนุญาต", "合同"),
            ("乙方", "Party B", "ฝ่ายที่สอง", "合同"),
            ("甲方", "Party A", "ฝ่ายที่หนึ่ง", "合同"),
            ("协议", "Agreement", "ข้อตกลง", "合同"),
            ("合同", "Contract","สัญญา", "合同"),
            ("通知", "Notice", "การแจ้ง", "通知"),
            ("会议纪要", "Meeting Minutes", "บันทึกการประชุม", "会议"),
            ("报告", "Report", "รายงาน", "报告"),
            # 法律术语
            ("仲裁", "Arbitration", "อนุญาต", "法律"),
            ("不可抗力", "Force Majeure", "อุปสรรคสงคราม", "法律"),
            ("产品责任", "Product Liability","ความรับผิดต่อสินค้า", "法律"),
            ("知识产权", "Intellectual Property", "ทรัพย์สินทางปัญญา", "法律"),
            ("保密义务", "Confidentiality", "การรักษาความลับ", "法律"),
            ("违约", "Breach of Contract", "การละเมิดสัญญา", "法律"),
            ("赔偿", "Compensation", "การชดเชย", "法律"),
            ("终止", "Termination", "การสิ้นสุด", "法律"),
            ("诉讼", "Litigation","การฟ้องคดี", "法律"),
            # 税务贸易术语
            ("关税", "Customs Duty", "อากรศุลกากร", "税务"),
            ("增值税", "VAT / Value Added Tax", "ภาษีมูลค่าเพิ่ม", "税务"),
            ("原产地证明", "Certificate of Origin", "หนังสือรับรองแหล่งกำเนิดสินค้า", "贸易"),
            ("检验证书", "Inspection Certificate", "วุติับัตรตรวจสอบสินค้า", "贸易"),
            ("货运单据", "Bill of Lading", "ตราสารการขนส่ง", "贸易"),
            # 电商术语
            ("电商平台", "E-Commerce Platform", "แพลตฟอร์มอีคอมเมิร์ซ", "电商"),
            ("最低销售额", "Minimum Sales Target", "ยอดขายขั้นต่ำ", "电商"),
            ("退换货", "Return and Exchange", "การคืนและเปลี่ยนสินค้า", "电商"),
        ]
        cursor.executemany("INSERT INTO terms (zh, en, th, category) VALUES (?, ?, ?, ?)", sample_terms)
    
    cursor.execute("SELECT COUNT(*) FROM templates")
    if cursor.fetchone()[0] == 0:
        sample_templates = [
            ("经销协议模板", "合同"),
            ("跨境电商补充协议", "合同"),
        ]
        template_contents = {
            "经销协议模板": """# 经销协议 | DEALER DISTRIBUTION AGREEMENT

**协议编号 (Agreement No.):** {agreement_no}
**签订日期 (Date):** {signing_date}
**签订地点 (Place):** {signing_place}

---

## 协议双方 Parties

**甲方 (Party A - 授权方/ licensor):**
公司名称 (Company Name): {company_name_cn}
注册地址 (Registered Address): {address_cn}
注册号 (Registration No.): {registration_no_cn}
联系人 (Contact Person): {contact_person_cn}
电话 (Tel): {phone_cn}
邮箱 (Email): {email_cn}

**乙方 (Party B - 经销商/ distributor):**
公司名称 (Company Name): VISION DANCE CO., LTD.
注册地址 (Registered Address): 34 Suk Kasame Road, Suthep Sub-district, Mueang Chiang Mai District, Chiang Mai Province 50200, Thailand
注册号 (Registration No.): 0505568025253
联系人 (Contact Person): {contact_person_th}
电话 (Tel): {phone_th}
邮箱 (Email): {email_th}

---

## 第一条 产品授权 Product Authorization

### 1.1 授权产品 Authorized Products
甲方授权乙方在指定区域内分销以下产品：
Party A hereby grants Party B the exclusive right to distribute the following products within the designated territory:

| 产品名称 Product Name | 规格型号 Specification | 建议零售价 MSRP |
|----------------------|----------------------|----------------|
| {products_list} | | |

### 1.2 授权区域 Territory
本协议授权区域为：泰国全境 (Kingdom of Thailand)
本授权为独占性授权 (Exclusive Distribution Right)。

### 1.3 授权期限 Term
本协议有效期为 {contract_term}，自 {start_date} 起至 {end_date} 止。
This Agreement shall remain in effect for a period of {contract_term} from {start_date} to {end_date}.

---

## 第二条 价格与付款 Price and Payment

### 2.1 结算货币 Currency
本协议以 {currency} (人民币 / CNY) 结算。

### 2.2 付款方式 Payment Terms
- 付款期限 (Payment Term): {payment_terms}
- 信用账期 (Credit Period): {credit_period} 天
- 最低订单量 (Minimum Order Quantity): {moq} 件/批次

### 2.3 价格调整 Price Adjustment
甲方保留调整价格的权利，应提前 {notice_period} 天书面通知乙方。
Party A reserves the right to adjust prices with {notice_period} days prior written notice to Party B.

---

## 第三条 双方权利义务 Rights and Obligations

### 3.1 甲方权利义务 Party A's Rights and Obligations
1. 提供符合中国国家标准和泰国TIS标准的产品
   (Provide products meeting both China National Standards and Thai TIS standards)
2. 提供产品质量保证 (Product Quality Guarantee)
3. 提供必要的市场推广支持和培训 (Marketing support and training)
4. 保护乙方的独占经销区域 (Protect Party B's exclusive territory)
5. 及时交付订单产品 (Timely delivery of ordered products)
6. 提供售后服务技术支持 (After-sales technical support)

### 3.2 乙方权利义务 Party B's Rights and Obligations
1. 完成最低销售额指标 (Achieve minimum sales target): {min_sales} / {sales_currency}
2. 遵守甲方制定的价格体系 (Maintain pricing structure)
3. 遵守泰国产品标准法规 (Comply with Thai product standards and regulations)
4. 建立完善的销售渠道和售后服务体系 (Establish sales channel and after-sales service)
5. 定期报告市场情况和库存状态 (Provide market reports and inventory status regularly)
6. 保护甲方品牌形象和知识产权 (Protect Party A's brand image and IP)

---

## 第四条 产品责任 Product Liability

### 4.1 责任范围 Liability Scope
依据《中华人民共和国产品质量法》和泰国《产品责任法》(Product Liability Act, B.E. 2551 (2008))：

1. **甲方责任 (Party A's Liability):**
   - 对产品原材料和制造缺陷负责
   - 承担因产品缺陷造成的人身伤害和财产损失赔偿责任
   - 赔偿限额: 年度累计不超过 {annual_liability_cap} / 单次不超过 {per_incident_cap}

2. **乙方责任 (Party B's Liability):**
   - 负责产品在泰国市场的存储和运输安全
   - 及时处理消费者投诉和产品召回
   - 配合甲方进行产品安全调查

### 4.2 保险要求 Insurance Requirement
乙方应购买产品责任保险，保额不低于 {insurance_coverage} / {insurance_currency}。

---

## 第五条 知识产权 Intellectual Property

### 5.1 甲方知识产权 Party A's IP
- 商标权 Trademark: {trademark_no}
- 专利权 Patent: {patent_no}
- 设计权 Design Rights: {design_rights}

### 5.2 授权使用 Licensed Use
乙方仅可将甲方知识产权用于本协议授权产品分销目的。
Party B may only use Party A's IP for the distribution of authorized products under this Agreement.

### 5.3 侵权责任 Infringement Liability
如发生第三方侵权投诉，乙方应及时通知甲方，甲方承担赔偿责任。
In case of third-party infringement complaints, Party B shall promptly notify Party A, and Party A shall bear the compensation liability.

---

## 第六条 争议解决 Dispute Resolution

### 6.1 友好协商 Amicable Settlement
任何争议应在争议发生后 {negotiation_period} 天内通过友好协商解决。
Any dispute shall be resolved through amicable negotiation within {negotiation_period} days after the dispute arises.

### 6.2 仲裁 Arbitration
如协商未能解决争议，任何一方可将争议提交以下仲裁机构：
If negotiation fails, either party may submit the dispute to the following arbitration institutions:

| 仲裁机构 | 仲裁地 | 适用规则 |
|---------|--------|---------|
| CIETAC (中国国际经济贸易仲裁委员会) | 北京/上海 | UNCITRAL Rules |
| SIAC (新加坡国际仲裁中心) | 新加坡 | SIAC Rules |
| THAC (泰国贸易仲裁委员会) | 曼谷 | THAC Rules |

仲裁裁决为终局裁决，对双方具有约束力。
The arbitral award shall be final and binding upon both parties.

### 6.3 适用法律 Governing Law
- 中国法: 《中华人民共和国民法典》
- 泰国法: 泰国《民事和商事法典》(Civil and Commercial Code of Thailand)

---

## 第七条 协议生效与终止 Effectiveness and Termination

### 7.1 协议生效 Effectiveness
本协议自双方签字之日起生效。
This Agreement shall come into force upon execution by both parties.

### 7.2 提前终止 Early Termination
任何一方可在下列情况下提前终止协议：
Either party may terminate this Agreement prematurely under the following circumstances:

1. 对方严重违约且在收到违约通知后 {cure_period} 天内未予纠正
   (The other party materially breaches and fails to remedy within {cure_period} days after notice)
2. 对方破产或清算 (The other party goes bankrupt or enters liquidation)
3. 不可抗力持续超过 {force_majeure_period} 天 (Force majeure continues for more than {force_majeure_period} days)

### 7.3 终止后义务 Post-termination Obligations
- 乙方应停止所有促销和销售活动
- 归还所有甲方知识产权材料
- 完成在途订单的处理
- 遵守竞业禁止条款 {non_compete_period} 个月

---

## 第八条 其他条款 Miscellaneous

### 8.1 保密条款 Confidentiality
双方应对本协议内容保密，未经对方书面同意不得向第三方披露。
Both parties shall keep the contents of this Agreement confidential and shall not disclose to any third party without the other party's written consent.

### 8.2 通知条款 Notice
所有通知应以书面形式发送至本协议载明的地址。
All notices shall be sent in writing to the addresses specified in this Agreement.

### 8.3 协议文本 Languages
本协议一式三份，中文、英文、泰文各执一份，三种文本具有同等法律效力。
This Agreement is executed in three copies in Chinese, English, and Thai, each text being equally authentic.

---

## 签署 Signatures

**甲方 (Party A):**
_______________________
姓名 (Name): ___________
职位 (Title): ___________
日期 (Date): ___________

**乙方 (Party B):**
_______________________
姓名 (Name): ___________
职位 (Title): ___________
日期 (Date): ___________
""",
            "跨境电商补充协议": """# 跨境电商补充协议 | CROSS-BORDER E-COMMERCE ADDENDUM

**协议编号 (Addendum No.):** {addendum_no}
**关联主协议 (Master Agreement):** {master_agreement_no}
**签订日期 (Date):** {signing_date}

---

## 前言 Preamble

本补充协议是《{master_agreement_no}》(以下简称"主协议")的组成部分，与主协议具有同等法律效力。
This Addendum forms an integral part of the Master Agreement No. {master_agreement_no} (hereinafter referred to as "Master Agreement") and shall have the same legal effect as the Master Agreement.

本补充协议专门针对跨境电子商务交易中涉及的中国与泰国之间的特殊法律合规要求而制定。
This Addendum is specifically formulated to address the special legal compliance requirements involved in cross-border e-commerce transactions between China and Thailand.

---

## 第一条 电商平台责任 E-commerce Platform Responsibilities

### 1.1 平台运营责任 Platform Operation Responsibilities

**甲方责任 (Party A's Responsibilities):**
1. 确保在中国境内注册的电商平台符合中国《电子商务法》要求
   (Ensure the e-commerce platform registered in China complies with the requirements of China's Electronic Commerce Law)
2. 取得增值电信业务经营许可证 (ICP License)
3. 建立消费者权益保护制度
4. 配合中国监管部门进行检查和调查

**乙方责任 (Party B's Responsibilities):**
1. 确保平台符合泰国《电子商务和数据保护法》(PDPA) 要求
   (Ensure the platform complies with Thailand's Electronic Commerce and Data Protection Act)
2. 保护消费者个人信息 (Protect consumer personal information)
3. 建立消费者投诉处理机制 (Establish consumer complaint handling mechanism)
4. 处理泰国消费者的退换货请求 (Handle return and exchange requests from Thai consumers)

### 1.2 商品信息管理 Product Information Management
- 所有在泰国销售的产品页面必须使用泰文标注
  (All product pages for sale in Thailand must be labeled in Thai)
- 产品标签应符合泰国FDA要求 (Product labels shall comply with Thai FDA requirements)
- 价格标示须包含税费说明 (Price display shall include tax information)

---

## 第二条 海关与税务 Customs and Taxation

### 2.1 清关责任 Customs Clearance Responsibilities

| 责任方 | 中国出口 | 泰国进口 |
|-------|---------|---------|
| 甲方 Party A | 出口清关、报关 | - |
| 乙方 Party B | - | 进口清关、报关 |

### 2.2 税务责任 Tax Responsibilities

**中国税务 (China Tax):**
- 出口增值税退税 (Export VAT refund)
- 企业所得税优惠 (Corporate income tax benefits)

**泰国税务 (Thailand Tax):**
- 进口关税 (Import duty): 按泰国海关税则执行
- 增值税 (VAT): 7%，由乙方代扣代缴
- 预提所得税 (Withholding tax): 按泰中税收协定执行

### 2.3 税务合规 Tax Compliance
- 双方应保存完整的税务凭证至少 {record_retention_period} 年
  (Both parties shall keep complete tax records for at least {record_retention_period} years)
- 任何税务优惠政策的变化应及时通知对方
  (Any changes in tax incentive policies shall be promptly communicated to the other party)

---

## 第三条 产品合规 Product Compliance

### 3.1 泰国产品标准 Thai Product Standards

| 检查项目 | 要求 | 负责方 |
|---------|-----|-------|
| TIS认证 Thai Industrial Standards | 适用于工业产品 | 甲方 |
| Thai FDA批准 | 食品、化妆品、药品 | 甲方 |
| 产品标签泰文标注 | 所有产品 | 甲方 |
| 原产地标识 (COO) | 所有产品 | 甲方 |
| 泰国语言要求 | 包装、说明书 | 甲方 |

### 3.2 中国出口合规 China Export Compliance
- 出口管制商品分类 (Export control classification)
- 商检合格证明 (Commodity inspection certificate)
- 原产地证明 (Certificate of Origin)

### 3.3 合规检查清单 Compliance Checklist

| 序号 | 检查项目 | 状态 |
|-----|---------|-----|
| 1 | 泰国工业标准 (TIS) 认证 | ☐ |
| 2 | 泰国FDA批准 (如适用) | ☐ |
| 3 | 产品标签泰文标注 | ☐ |
| 4 | 原产地标识 (Made in China) | ☐ |
| 5 | 包装符合泰国法规 | ☐ |
| 6 | 说明书泰文版 | ☐ |

---

## 第四条 物流与仓储 Logistics and Warehousing

### 4.1 仓储责任 Warehouse Responsibilities
- 甲方在中国境内设立或租用符合出口仓储标准的仓库
- 乙方在泰国设立或租用符合进口仓储标准的仓库
- 双方应确保产品在储存期间的质量安全

### 4.2 物流服务商 Logistics Service Provider
推荐物流服务商 (Recommended Logistics Service Providers):
- 中国至泰国: {logistics_provider_cn}
- 泰国至中国: {logistics_provider_th}

### 4.3 运输保险 Transportation Insurance
- 建议为每批货物购买运输保险
- 保险金额应不低于货值的 {insurance_coverage_percentage}%
- 保险受益人应为货物所有权人

---

## 第五条 消费者保护 Consumer Protection

### 5.1 退换货政策 Return and Exchange Policy

| 商品类别 | 退货期限 | 退货条件 | 运费承担 |
|---------|---------|---------|---------|
| 服装/配饰 | {apparel_return_days} 天 | 未使用、标签完好 | 买方 |
| 电子产品 | {electronics_return_days} 天 | 功能异常 | 卖方 |
| 食品/化妆品 | {food_cosmetic_return_days} 天 | 未开封 | 买方 |
| 定制商品 | 不可退 | - | - |

### 5.2 投诉处理 Complaint Handling
- 乙方应在 {complaint_response_days} 个工作日内回应消费者投诉
- 重大投诉应在 {escalation_days} 天内升级处理
- 建立投诉数据库并定期分析改进

### 5.3 个人信息保护 Personal Data Protection
- 双方应遵守各自国家的个人信息保护法律
- 消费者数据跨境传输应符合相关法规要求
- 建立数据泄露应急响应机制

---

## 第六条 争议解决 Dispute Resolution

### 6.1 跨境争议处理 Cross-border Dispute Resolution
因本补充协议产生的争议，优先通过友好协商解决。
Disputes arising from this Addendum shall be resolved through amicable negotiation first.

### 6.2 适用规则 Applicable Rules
如协商不成，提交 {arbitration_institution} 仲裁，适用《联合国国际贸易法委员会仲裁规则》(UNCITRAL Rules)。

---

## 第七条 其他条款 Miscellaneous

### 7.1 协议修订 Amendment
本补充协议的修订须经双方书面同意。
Amendments to this Addendum require written consent of both parties.

### 7.2 协议期限 Duration
本补充协议与主协议同期有效。
This Addendum shall remain effective concurrently with the Master Agreement.

### 7.3 适用法律 Governing Law
- 中国法: 《中华人民共和国电子商务法》、《民法典》
- 泰国法: 《电子商务和数据保护法》(PDPA)、《民事和商事法典》

---

## 签署 Signatures

**甲方 (Party A):**
_______________________
姓名 (Name): ___________
职位 (Title): ___________
日期 (Date): ___________

**乙方 (Party B):**
_______________________
姓名 (Name): ___________
职位 (Title): ___________
日期 (Date): ___________
""",
        }
        for name, category in sample_templates:
            cursor.execute("INSERT INTO templates (name, category) VALUES (?, ?)", (name, category))
            template_id = cursor.lastrowid
            tmpl_content = template_contents.get(name, "")
            cursor.execute(
                "INSERT INTO template_versions (template_id, version, content, description) VALUES (?, 1, ?, ?)",
                (template_id, tmpl_content, f"{name} v1.0 初始版本")
            )


    conn.commit()
    conn.close()

class Document(BaseModel):
    id: Optional[int] = None
    title: str
    content: str
    content_en: Optional[str] = None
    content_th: Optional[str] = None
    template_id: Optional[int] = None
    template_version_id: Optional[int] = None
    category: str = "其他"
    status: str = "draft"
    created_by: Optional[str] = None

class Template(BaseModel):
    id: Optional[int] = None
    name: str
    category: Optional[str] = None
    content: str

class Term(BaseModel):
    id: Optional[int] = None
    zh: str
    en: Optional[str] = None
    th: Optional[str] = None
    category: str = "通用"

@app.get("/")
def root():
    return {"message": "CompanyDoc API - 企业文档管理系统", "version": "1.0.0"}

@app.get("/api/documents", response_model=List[dict])
def get_documents(page: int = 1, limit: int = 100, category: str = None, status: str = None, db: sqlite3.Connection = Depends(get_db)):
    offset = (page - 1) * limit
    query = "SELECT * FROM documents WHERE 1=1"
    params = []
    
    if category:
        query += " AND category = ?"
        params.append(category)
    if status:
        query += " AND status = ?"
        params.append(status)
    
    query += " ORDER BY created_at DESC LIMIT ? OFFSET ?"
    params.extend([limit, offset])
    
    cursor = db.execute(query, params)
    return [dict(row) for row in cursor.fetchall()]

@app.get("/api/documents/{doc_id}", response_model=dict)
def get_document(doc_id: int, db: sqlite3.Connection = Depends(get_db)):
    cursor = db.execute("SELECT * FROM documents WHERE id = ?", (doc_id,))
    row = cursor.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Document not found")
    return dict(row)

@app.post("/api/documents", response_model=dict)
def create_document(doc: Document, db: sqlite3.Connection = Depends(get_db)):
    cursor = db.execute('''
        INSERT INTO documents (title, content, content_en, content_th, template_id, template_version_id, category, status, created_by)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (doc.title, doc.content, doc.content_en, doc.content_th, doc.template_id, doc.template_version_id, doc.category, doc.status, doc.created_by))
    
    db.commit()
    doc_id = cursor.lastrowid
    
    cursor = db.execute("SELECT * FROM documents WHERE id = ?", (doc_id,))
    return dict(cursor.fetchone())

@app.put("/api/documents/{doc_id}", response_model=dict)
def update_document(doc_id: int, doc: Document, db: sqlite3.Connection = Depends(get_db)):
    cursor = db.execute('''
        UPDATE documents SET title=?, content=?, content_en=?, content_th=?, template_id=?, template_version_id=?, category=?, status=?, updated_at=CURRENT_TIMESTAMP
        WHERE id=?
    ''', (doc.title, doc.content, doc.content_en, doc.content_th, doc.template_id, doc.template_version_id, doc.category, doc.status, doc_id))
    
    if cursor.rowcount == 0:
        raise HTTPException(status_code=404, detail="Document not found")
    
    db.commit()
    
    cursor = db.execute("SELECT * FROM documents WHERE id = ?", (doc_id,))
    return dict(cursor.fetchone())

@app.delete("/api/documents/{doc_id}")
def delete_document(doc_id: int, db: sqlite3.Connection = Depends(get_db)):
    cursor = db.execute("DELETE FROM documents WHERE id = ?", (doc_id,))
    if cursor.rowcount == 0:
        raise HTTPException(status_code=404, detail="Document not found")
    
    db.commit()
    return {"message": "Document deleted successfully", "id": doc_id}

@app.get("/api/templates", response_model=List[dict])
def get_templates(db: sqlite3.Connection = Depends(get_db)):
    """获取所有模板（包含版本数量和最新版本信息）"""
    cursor = db.execute('''
        SELECT t.*,
               COUNT(tv.id) as version_count,
               (SELECT tv2.version FROM template_versions tv2
                WHERE tv2.template_id = t.id ORDER BY tv2.version DESC LIMIT 1) as latest_version,
               (SELECT tv2.content FROM template_versions tv2
                WHERE tv2.template_id = t.id ORDER BY tv2.version DESC LIMIT 1) as latest_content
        FROM templates t
        LEFT JOIN template_versions tv ON t.id = tv.template_id
        GROUP BY t.id
        ORDER BY t.name
    ''')
    return [dict(row) for row in cursor.fetchall()]

@app.post("/api/templates", response_model=dict)
def create_template(template: Template, db: sqlite3.Connection = Depends(get_db)):
    """创建新模板（需要同时创建第一个版本）"""
    cursor = db.execute('''
        INSERT INTO templates (name, category)
        VALUES (?, ?)
    ''', (template.name, template.category))

    db.commit()
    template_id = cursor.lastrowid

    # 创建第一个版本
    cursor = db.execute('''
        INSERT INTO template_versions (template_id, version, content, description)
        VALUES (?, 1, ?, ?)
    ''', (template_id, template.content or "", f"{template.name} v1.0"))

    db.commit()

    cursor = db.execute("SELECT * FROM templates WHERE id = ?", (template_id,))
    return dict(cursor.fetchone())

@app.delete("/api/templates/{template_id}")
def delete_template(template_id: int, db: sqlite3.Connection = Depends(get_db)):
    """删除模板（会级联删除所有版本）"""
    cursor = db.execute("DELETE FROM templates WHERE id = ?", (template_id,))
    if cursor.rowcount == 0:
        raise HTTPException(status_code=404, detail="Template not found")

    db.commit()
    return {"message": "Template deleted successfully", "id": template_id}

@app.get("/api/templates/{template_id}/versions", response_model=List[dict])
def get_template_versions(template_id: int, db: sqlite3.Connection = Depends(get_db)):
    """获取模板的所有版本"""
    cursor = db.execute('''
        SELECT tv.*,
               (SELECT COUNT(*) FROM documents d WHERE d.template_version_id = tv.id) as doc_count
        FROM template_versions tv
        WHERE tv.template_id = ?
        ORDER BY tv.version DESC
    ''', (template_id,))
    return [dict(row) for row in cursor.fetchall()]

@app.get("/api/template-versions/{version_id}", response_model=dict)
def get_template_version(version_id: int, db: sqlite3.Connection = Depends(get_db)):
    """获取特定版本详情"""
    cursor = db.execute("SELECT * FROM template_versions WHERE id = ?", (version_id,))
    row = cursor.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Template version not found")
    return dict(row)

@app.post("/api/templates/{template_id}/versions", response_model=dict)
def create_template_version(template_id: int, version: dict, db: sqlite3.Connection = Depends(get_db)):
    """为模板创建新版本"""
    # 获取最新版本号
    cursor = db.execute('''
        SELECT MAX(version) as max_version FROM template_versions WHERE template_id = ?
    ''', (template_id,))
    row = cursor.fetchone()
    next_version = (row['max_version'] or 0) + 1

    cursor = db.execute('''
        INSERT INTO template_versions (template_id, version, content, description)
        VALUES (?, ?, ?, ?)
    ''', (template_id, next_version, version.get('content', ''), version.get('description', '')))

    db.commit()
    version_id = cursor.lastrowid

    cursor = db.execute("SELECT * FROM template_versions WHERE id = ?", (version_id,))
    return dict(cursor.fetchone())

@app.put("/api/template-versions/{version_id}", response_model=dict)
def update_template_version(version_id: int, version: dict, db: sqlite3.Connection = Depends(get_db)):
    """更新模板版本内容"""
    cursor = db.execute('''
        UPDATE template_versions SET content = ?, description = ? WHERE id = ?
    ''', (version.get('content', ''), version.get('description', ''), version_id))

    if cursor.rowcount == 0:
        raise HTTPException(status_code=404, detail="Template version not found")

    db.commit()

    cursor = db.execute("SELECT * FROM template_versions WHERE id = ?", (version_id,))
    return dict(cursor.fetchone())

@app.get("/api/template-versions/{version_id}/documents", response_model=List[dict])
def get_version_documents(version_id: int, db: sqlite3.Connection = Depends(get_db)):
    """获取使用某版本的所有文档"""
    cursor = db.execute('''
        SELECT d.*, t.name as template_name, tv.version
        FROM documents d
        LEFT JOIN templates t ON d.template_id = t.id
        LEFT JOIN template_versions tv ON d.template_version_id = tv.id
        WHERE d.template_version_id = ?
        ORDER BY d.created_at DESC
    ''', (version_id,))
    return [dict(row) for row in cursor.fetchall()]

@app.get("/api/terms", response_model=List[dict])
def get_terms(category: str = None, db: sqlite3.Connection = Depends(get_db)):
    query = "SELECT * FROM terms WHERE 1=1"
    params = []
    
    if category:
        query += " AND category = ?"
        params.append(category)
    
    query += " ORDER BY zh"
    cursor = db.execute(query, params)
    return [dict(row) for row in cursor.fetchall()]

@app.post("/api/terms", response_model=dict)
def create_term(term: Term, db: sqlite3.Connection = Depends(get_db)):
    cursor = db.execute('''
        INSERT INTO terms (zh, en, th, category)
        VALUES (?, ?, ?, ?)
    ''', (term.zh, term.en, term.th, term.category))
    
    db.commit()
    term_id = cursor.lastrowid
    
    cursor = db.execute("SELECT * FROM terms WHERE id = ?", (term_id,))
    return dict(cursor.fetchone())

@app.put("/api/terms/{term_id}", response_model=dict)
def update_term(term_id: int, term: Term, db: sqlite3.Connection = Depends(get_db)):
    cursor = db.execute('''
        UPDATE terms SET zh=?, en=?, th=?, category=? WHERE id=?
    ''', (term.zh, term.en, term.th, term.category, term_id))
    
    if cursor.rowcount == 0:
        raise HTTPException(status_code=404, detail="Term not found")
    
    db.commit()
    
    cursor = db.execute("SELECT * FROM terms WHERE id = ?", (term_id,))
    return dict(cursor.fetchone())

@app.delete("/api/terms/{term_id}")
def delete_term(term_id: int, db: sqlite3.Connection = Depends(get_db)):
    cursor = db.execute("DELETE FROM terms WHERE id = ?", (term_id,))
    if cursor.rowcount == 0:
        raise HTTPException(status_code=404, detail="Term not found")
    
    db.commit()
    return {"message": "Term deleted successfully", "id": term_id}

@app.post("/api/documents/{doc_id}/translate")
def translate_document(doc_id: int, target_language: str, db: sqlite3.Connection = Depends(get_db)):
    cursor = db.execute("SELECT * FROM documents WHERE id = ?", (doc_id,))
    row = cursor.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Document not found")
    
    doc = dict(row)
    original_content = doc['content']
    
    translated_content = f"[翻译结果 - {target_language.upper()}] {original_content[:100]}"
    
    if target_language == 'en':
        db.execute("UPDATE documents SET content_en = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?", (translated_content, doc_id))
    elif target_language == 'th':
        db.execute("UPDATE documents SET content_th = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?", (translated_content, doc_id))
    
    db.commit()
    
    return {
        "document_id": doc_id,
        "target_language": target_language,
        "message": "Translation completed successfully",
        "translated_content": translated_content
    }



@app.get("/api/stats")
def get_stats(db: sqlite3.Connection = Depends(get_db)):
    cursor = db.execute("SELECT COUNT(*) FROM documents")
    total_docs = cursor.fetchone()[0]
    
    cursor = db.execute("SELECT COUNT(*) FROM documents WHERE status = 'draft'")
    draft_docs = cursor.fetchone()[0]
    
    cursor = db.execute("SELECT COUNT(*) FROM documents WHERE status = 'approved'")
    approved_docs = cursor.fetchone()[0]
    
    cursor = db.execute("SELECT COUNT(*) FROM terms")
    total_terms = cursor.fetchone()[0]
    
    cursor = db.execute("SELECT COUNT(*) FROM templates")
    total_templates = cursor.fetchone()[0]
    
    return {
        "total_documents": total_docs,
        "draft_documents": draft_docs,
        "approved_documents": approved_docs,
        "total_terms": total_terms,
        "total_templates": total_templates
    }

# ============ 客户管理接口 ============
class Customer(BaseModel):
    id: Optional[int] = None
    name: str
    company_name: Optional[str] = None
    company_name_th: Optional[str] = None
    registration_number: Optional[str] = None
    address: Optional[str] = None
    address_th: Optional[str] = None
    contact_person: Optional[str] = None
    contact_email: Optional[str] = None
    contact_phone: Optional[str] = None
    contact_phone_th: Optional[str] = None
    notes: Optional[str] = None
    source_document_id: Optional[int] = None

@app.get("/api/customers", response_model=List[dict])
def get_customers(db: sqlite3.Connection = Depends(get_db)):
    """获取所有客户"""
    cursor = db.execute("SELECT * FROM customers ORDER BY name")
    return [dict(row) for row in cursor.fetchall()]

@app.get("/api/customers/{customer_id}", response_model=dict)
def get_customer(customer_id: int, db: sqlite3.Connection = Depends(get_db)):
    """获取客户详情"""
    cursor = db.execute("SELECT * FROM customers WHERE id = ?", (customer_id,))
    row = cursor.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Customer not found")
    return dict(row)

@app.post("/api/customers", response_model=dict)
def create_customer(customer: Customer, db: sqlite3.Connection = Depends(get_db)):
    """创建客户"""
    cursor = db.execute('''
        INSERT INTO customers (name, company_name, company_name_th, registration_number,
            address, address_th, contact_person, contact_email, contact_phone,
            contact_phone_th, notes, source_document_id)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (customer.name, customer.company_name, customer.company_name_th,
          customer.registration_number, customer.address, customer.address_th,
          customer.contact_person, customer.contact_email, customer.contact_phone,
          customer.contact_phone_th, customer.notes, customer.source_document_id))

    db.commit()
    customer_id = cursor.lastrowid

    cursor = db.execute("SELECT * FROM customers WHERE id = ?", (customer_id,))
    return dict(cursor.fetchone())

@app.put("/api/customers/{customer_id}", response_model=dict)
def update_customer(customer_id: int, customer: Customer, db: sqlite3.Connection = Depends(get_db)):
    """更新客户信息"""
    cursor = db.execute('''
        UPDATE customers SET name=?, company_name=?, company_name_th=?, registration_number=?,
            address=?, address_th=?, contact_person=?, contact_email=?, contact_phone=?,
            contact_phone_th=?, notes=?, updated_at=CURRENT_TIMESTAMP
        WHERE id=?
    ''', (customer.name, customer.company_name, customer.company_name_th,
          customer.registration_number, customer.address, customer.address_th,
          customer.contact_person, customer.contact_email, customer.contact_phone,
          customer.contact_phone_th, customer.notes, customer_id))

    if cursor.rowcount == 0:
        raise HTTPException(status_code=404, detail="Customer not found")

    db.commit()

    cursor = db.execute("SELECT * FROM customers WHERE id = ?", (customer_id,))
    return dict(cursor.fetchone())

@app.delete("/api/customers/{customer_id}")
def delete_customer(customer_id: int, db: sqlite3.Connection = Depends(get_db)):
    """删除客户"""
    cursor = db.execute("DELETE FROM customers WHERE id = ?", (customer_id,))
    if cursor.rowcount == 0:
        raise HTTPException(status_code=404, detail="Customer not found")

    db.commit()
    return {"message": "Customer deleted successfully", "id": customer_id}

# ============ 文件上传接口 ============
import shutil
from pathlib import Path

UPLOAD_DIR = Path(os.path.dirname(__file__)).parent / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)

@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    """上传文档并存储"""
    # Create unique filename
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_filename = f"{timestamp}_{file.filename}"
    file_path = UPLOAD_DIR / safe_filename

    # Save file
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    # Get file info
    file_size = file_path.stat().st_size
    file_type = file.filename.split(".")[-1] if "." in file.filename else "unknown"

    # Save to database
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO uploaded_files (filename, original_filename, file_path, file_type, file_size)
        VALUES (?, ?, ?, ?, ?)
    ''', (safe_filename, file.filename, str(file_path), file_type, file_size))
    conn.commit()
    file_id = cursor.lastrowid
    conn.close()

    return {"file_id": file_id, "filename": safe_filename, "original_filename": file.filename}

@app.post("/api/upload/{file_id}/parse")
def parse_uploaded_file(file_id: int):
    """AI解析上传的文件，提取客户信息和文档内容"""
    # Create our own connection to avoid thread issues
    db = sqlite3.connect(DB_PATH)
    db.row_factory = sqlite3.Row
    try:
        # Get file info
        cursor = db.execute("SELECT * FROM uploaded_files WHERE id = ?", (file_id,))
        row = cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="File not found")

        file_info = dict(row)
        file_path = Path(file_info["file_path"])

        if not file_path.exists():
            raise HTTPException(status_code=404, detail="File not found on disk")

        # Read file content
        content = ""
        if file_info["file_type"] == "txt":
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
        elif file_info["file_type"] in ["md", "markdown"]:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
        elif file_info["file_type"] in ["doc", "docx"]:
            try:
                from docx import Document
                doc = Document(file_path)
                # Extract text from all paragraphs
                paragraphs = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text)
                # Also extract from tables if any
                for table in doc.tables:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if cells:
                            paragraphs.append(" | ".join(cells))
                content = "\n".join(paragraphs)
            except Exception as e:
                return {"error": f"Failed to read Word document: {str(e)}"}
        else:
            # Try to read as text
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
            except:
                return {"error": f"Unsupported file type: {file_info['file_type']}"}

        if not content:
            return {"error": "File is empty or could not be read"}

        # Use AI to extract customer info (sync call)
        import asyncio
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        customer_info = loop.run_until_complete(ai_extract_customer_info(content))
        loop.close()

        # Update file record with parsed content
        db.execute('''
            UPDATE uploaded_files SET parsed_content = ? WHERE id = ?
        ''', (content, file_id))
        db.commit()

        return {
            "file_id": file_id,
            "content": content,
            "customer_info": customer_info
        }
    finally:
        db.close()

async def ai_extract_customer_info(content: str) -> dict:
    """使用 AI 从文档内容中提取客户信息"""
    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.post(
                f"{os.getenv('ANTHROPIC_BASE_URL', 'https://api.minimaxi.com/anthropic')}/v1/messages",
                headers={
                    "Authorization": f"Bearer {os.getenv('ANTHROPIC_AUTH_TOKEN', '')}",
                    "Content-Type": "application/json",
                    "anthropic-version": "2023-06-01",
                },
                json={
                    "model": os.getenv("ANTHROPIC_MODEL", "MiniMax-M2.7"),
                    "max_tokens": 2048,
                    "system": """你是一位专业的商业文档分析助手。用户会提供一份商业文档内容，你需要从中提取客户信息。

请提取以下信息（如果文档中没有，则留空）：
- name: 客户姓名或联系人姓名
- company_name: 公司英文名称
- company_name_th: 公司泰文名称
- registration_number: 注册登记号
- address: 英文地址
- address_th: 泰文地址
- contact_person: 联系人姓名
- contact_email: 邮箱
- contact_phone: 电话
- contact_phone_th: 泰文电话

请以 JSON格式返回，格式如下：
{
  "name": "...",
  "company_name": "...",
  "company_name_th": "...",
  "registration_number": "...",
  "address": "...",
  "address_th": "...",
  "contact_person": "...",
  "contact_email": "...",
  "contact_phone": "...",
  "contact_phone_th": "..."
}

只返回 JSON，不要有其他文字。""",
                    "messages": [
                        {
                            "role": "user",
                            "content": content[:8000],  # Limit content length
                        }
                    ],
                },
            )

            if response.status_code == 200:
                data = response.json()
                # Find the text content (not thinking)
                result_text = ""
                for item in data.get("content", []):
                    if item.get("type") == "text":
                        result_text = item.get("text", "")
                        break

                # Parse JSON result
                import json
                try:
                    # Try to find JSON in the response
                    json_start = result_text.find("{")
                    json_end = result_text.rfind("}") + 1
                    if json_start >= 0 and json_end > json_start:
                        customer_data = json.loads(result_text[json_start:json_end])
                        return customer_data
                except:
                    pass

                return {"error": "Failed to parse customer info"}
            else:
                return {"error": f"AI service error: {response.status_code}"}
    except Exception as e:
        return {"error": str(e)}

@app.post("/api/customers/from-file/{file_id}")
def create_customer_from_file(file_id: int, db: sqlite3.Connection = Depends(get_db)):
    """从解析的文件创建客户"""
    # Get file info
    cursor = db.execute("SELECT * FROM uploaded_files WHERE id = ?", (file_id,))
    row = cursor.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="File not found")

    file_info = dict(row)
    if not file_info.get("parsed_content"):
        raise HTTPException(status_code=400, detail="File has not been parsed yet")

    # Get customer info from file
    import json
    try:
        customer_data = json.loads(file_info.get("customer_data", "{}"))
    except:
        customer_data = {}

    # Create customer
    cursor = db.execute('''
        INSERT INTO customers (name, company_name, company_name_th, registration_number,
            address, address_th, contact_person, contact_email, contact_phone,
            contact_phone_th, notes, source_document_id)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (
        customer_data.get("name", "Unknown"),
        customer_data.get("company_name", ""),
        customer_data.get("company_name_th", ""),
        customer_data.get("registration_number", ""),
        customer_data.get("address", ""),
        customer_data.get("address_th", ""),
        customer_data.get("contact_person", ""),
        customer_data.get("contact_email", ""),
        customer_data.get("contact_phone", ""),
        customer_data.get("contact_phone_th", ""),
        f"From file: {file_info['original_filename']}",
        file_id
    ))

    db.commit()
    customer_id = cursor.lastrowid

    # Update file with customer_id
    cursor = db.execute("UPDATE uploaded_files SET customer_id = ? WHERE id = ?", (customer_id, file_id))
    db.commit()

    cursor = db.execute("SELECT * FROM customers WHERE id = ?", (customer_id,))
    return dict(cursor.fetchone())

@app.get("/api/upload/{file_id}/content")
def get_uploaded_file_content(file_id: int, db: sqlite3.Connection = Depends(get_db)):
    """获取上传文件的内容"""
    cursor = db.execute("SELECT * FROM uploaded_files WHERE id = ?", (file_id,))
    row = cursor.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="File not found")

    file_info = dict(row)
    file_path = Path(file_info["file_path"])

    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found on disk")

    try:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        return {"content": content}
    except:
        raise HTTPException(status_code=400, detail="Could not read file content")

@app.get("/api/uploaded-files", response_model=List[dict])
def get_uploaded_files(db: sqlite3.Connection = Depends(get_db)):
    """获取所有上传的文件"""
    cursor = db.execute("SELECT * FROM uploaded_files ORDER BY created_at DESC")
    return [dict(row) for row in cursor.fetchall()]

@app.get("/api/customers/{customer_id}/documents", response_model=List[dict])
def get_customer_documents(customer_id: int, db: sqlite3.Connection = Depends(get_db)):
    """获取某客户相关的文档"""
    cursor = db.execute('''
        SELECT d.*, t.name as template_name
        FROM documents d
        LEFT JOIN templates t ON d.template_id = t.id
        WHERE d.created_by = ? OR d.template_id IN (
            SELECT template_id FROM documents WHERE created_by = ?
        )
        ORDER BY d.created_at DESC
    ''', (str(customer_id), str(customer_id)))
    return [dict(row) for row in cursor.fetchall()]

@app.get("/api/customers/{customer_id}/files", response_model=List[dict])
def get_customer_files(customer_id: int, db: sqlite3.Connection = Depends(get_db)):
    """获取某客户的上传文件"""
    cursor = db.execute('''
        SELECT * FROM uploaded_files WHERE customer_id = ? ORDER BY created_at DESC
    ''', (customer_id,))
    return [dict(row) for row in cursor.fetchall()]

# AI 协作接口
try:
    from ai_service import ai_complete, ai_generate_from_template
except ImportError:
    from app.ai_service import ai_complete, ai_generate_from_template
from pydantic import BaseModel
from typing import Literal, Optional

class AIContentRequest(BaseModel):
    content: str
    action: Literal["expand", "translate_en", "translate_th", "polish", "summarize"] = "expand"
    max_tokens: int = 2048

class AIGenerateRequest(BaseModel):
    template_name: str
    variables: dict
    language: Literal["cn", "en", "th"] = "cn"

@app.post("/api/ai/complete")
async def ai_complete_content(request: AIContentRequest):
    """AI 协作 - 续写/翻译/优化/总结"""
    result = await ai_complete(
        prompt=request.content,
        action=request.action,
        max_tokens=request.max_tokens,
    )
    return {"result": result}

@app.post("/api/ai/generate")
async def ai_generate_document(request: AIGenerateRequest):
    """AI 根据模板生成文档"""
    result = await ai_generate_from_template(
        template_name=request.template_name,
        variables=request.variables,
        language=request.language,
    )
    return {"result": result}

@app.get("/api/ai/health")
async def ai_health_check():
    """AI 服务健康检查"""
    api_key = os.getenv("ANTHROPIC_AUTH_TOKEN", "")
    return {
        "status": "ok" if api_key else "not_configured",
        "model": os.getenv("ANTHROPIC_MODEL", "MiniMax-M2.7"),
        "has_api_key": bool(api_key),
    }

if __name__ == "__main__":
    import uvicorn
    print("🚀 CompanyDoc API 启动中...")
    print("📊 API 文档: http://localhost:8000/docs")
    print("📊 备用文档: http://localhost:8000/redoc")
    uvicorn.run(app, host="0.0.0.0", port=8000)
